"""
النموذج الأول: python-docx-template
توليد وثيقة Word عربية من قالب بمتغيرات قابلة للاستبدال
الفائدة: إنتاج تقارير الجرد الوطني والمذكرات الرسمية من قالب واحد
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ========== بيانات المتغيرات ==========
# هذه هي البيانات التي تتغير في كل تقرير
data = {
    "عنوان_التقرير": "تقرير الجرد الوطني لغازات الاحتباس الحراري",
    "السنة": "2024",
    "القطاع": "قطاع الطاقة",
    "اجمالي_الانبعاثات": "١٨٥.٤",
    "الوحدة": "مليون طن ثاني أكسيد الكربون المكافئ",
    "المقارنة": "زيادة بنسبة ٣.٢٪ مقارنة بعام ٢٠٢٣",
    "اسم_المحرر": "قسم سياسات المناخ",
    "التاريخ": "مايو ٢٠٢٥",
    "بيانات_الجدول": [
        {"الفئة": "احتراق الوقود", "الانبعاثات": "١٤٢.٣", "النسبة": "٧٦.٨٪"},
        {"الفئة": "العمليات الصناعية", "الانبعاثات": "٢٨.١", "النسبة": "١٥.٢٪"},
        {"الفئة": "التسرب", "الانبعاثات": "١٥.٠", "النسبة": "٨.١٪"},
    ]
}

# ========== دالة مساعدة: تعيين RTL ==========
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_cell_rtl(cell):
    for para in cell.paragraphs:
        set_rtl(para)
        for run in para.runs:
            run.font.name = "Simplified Arabic"
            run.font.size = Pt(12)

# ========== بناء الوثيقة ==========
doc = Document()

# إعداد الخط الافتراضي
style = doc.styles['Normal']
style.font.name = "Simplified Arabic"
style.font.size = Pt(12)

# هوامش الصفحة
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ===== العنوان الرئيسي =====
title = doc.add_paragraph()
set_rtl(title)
title.paragraph_format.space_after = Pt(6)
run = title.add_run(data["عنوان_التقرير"])
run.bold = True
run.font.name = "Simplified Arabic"
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ===== معلومات التقرير =====
info = doc.add_paragraph()
set_rtl(info)
r = info.add_run(f"السنة المرجعية: {data['السنة']}    |    القطاع: {data['القطاع']}")
r.font.name = "Simplified Arabic"
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# ===== ملخص الانبعاثات =====
h1 = doc.add_paragraph()
set_rtl(h1)
rh = h1.add_run("ملخص الانبعاثات")
rh.bold = True
rh.font.name = "Simplified Arabic"
rh.font.size = Pt(14)
rh.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

body = doc.add_paragraph()
set_rtl(body)
rb = body.add_run(
    f"بلغ إجمالي انبعاثات {data['القطاع']} خلال عام {data['السنة']} "
    f"نحو {data['اجمالي_الانبعاثات']} {data['الوحدة']}، "
    f"وهو ما يمثل {data['المقارنة']}."
)
rb.font.name = "Simplified Arabic"
rb.font.size = Pt(12)

doc.add_paragraph()

# ===== جدول البيانات =====
h2 = doc.add_paragraph()
set_rtl(h2)
rh2 = h2.add_run("تفصيل الانبعاثات حسب الفئة")
rh2.bold = True
rh2.font.name = "Simplified Arabic"
rh2.font.size = Pt(14)
rh2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# رأس الجدول
headers = ["الفئة", "الانبعاثات (مليون طن)", "النسبة"]
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    set_cell_rtl(hdr_cells[i])
    for run in hdr_cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tc = hdr_cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1F497D')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

# صفوف البيانات
for row_data in data["بيانات_الجدول"]:
    row = table.add_row()
    row.cells[0].text = row_data["الفئة"]
    row.cells[1].text = row_data["الانبعاثات"]
    row.cells[2].text = row_data["النسبة"]
    for cell in row.cells:
        set_cell_rtl(cell)

doc.add_paragraph()

# ===== التذييل =====
footer_p = doc.add_paragraph()
set_rtl(footer_p)
rf = footer_p.add_run(f"أعده: {data['اسم_المحرر']}    |    التاريخ: {data['التاريخ']}")
rf.font.name = "Simplified Arabic"
rf.font.size = Pt(10)
rf.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
rf.italic = True

# ========== حفظ الملف ==========
output_path = "/home/claude/skills-repo/demos/output_demo1_تقرير_جرد.docx"
doc.save(output_path)
print(f"تم إنشاء الملف: {output_path}")
