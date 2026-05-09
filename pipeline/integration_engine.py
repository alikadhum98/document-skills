"""
محرك ربط الأدوات
يجمع مخرجات جميع الأدوات ويولّد:
1. لوحة تحكم شاملة (Word)
2. ملف JSON موحّد للحالة الكاملة
3. يُحدّث بيانات الأدوات من المصدر المشترك
"""

import sys, os, json, datetime, shutil
ROOT = os.path.dirname(os.path.dirname(__file__))
sys.path.insert(0, ROOT)
sys.path.insert(0, os.path.join(ROOT, 'core'))

from shared_data import (PARTY, EMISSIONS, LATEST_YEAR, BASE_YEAR,
                          NDC, FINANCE_MOBILIZED, ARTICLE6,
                          CURRENT_CONFERENCE, NDC_NAVIGATOR_SCORES,
                          SECTOR_SHARES_LATEST, LAST_UPDATED)

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import arabic_reshaper
from bidi.algorithm import get_display
import numpy as np
import io

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

plt.rcParams['font.family'] = 'DejaVu Sans'
def ar(t): return get_display(arabic_reshaper.reshape(str(t)))

C_DARK=RGBColor(0x1F,0x49,0x7D); C_MED=RGBColor(0x2E,0x75,0xB6)
C_GREEN=RGBColor(0x37,0x86,0x44); C_RED=RGBColor(0xC0,0x00,0x00)
C_AMBER=RGBColor(0xBF,0x8F,0x00); C_WHITE=RGBColor(0xFF,0xFF,0xFF)
C_GRAY=RGBColor(0x70,0x70,0x70); C_TEXT=RGBColor(0x26,0x26,0x26)
FONT="Simplified Arabic"

# =========================================================
# حسابات مشتقة من البيانات المشتركة
# =========================================================
def compute_all():
    total_mob    = FINANCE_MOBILIZED["إجمالي_مليون"] / 1000
    gap          = NDC["احتياج_تمويل_مليار"] - total_mob
    cov_pct      = total_mob / NDC["احتياج_تمويل_مليار"] * 100
    ndc_avg      = sum(NDC_NAVIGATOR_SCORES.values()) / len(NDC_NAVIGATOR_SCORES)
    chg_base     = (EMISSIONS[LATEST_YEAR] - EMISSIONS[BASE_YEAR]) / EMISSIONS[BASE_YEAR] * 100
    chg_yoy      = (EMISSIONS[LATEST_YEAR] - EMISSIONS[2020]) / EMISSIONS[2020] * 100
    ndc_gap      = EMISSIONS[LATEST_YEAR] - NDC["هدف_غير_مشروط_tCO2eq"] / 1000
    return {
        "تمويل_مُعبَّأ_مليار": round(total_mob, 3),
        "فجوة_تمويل_مليار":    round(gap, 2),
        "تغطية_pct":           round(cov_pct, 3),
        "متوسط_NDC_Navigator": round(ndc_avg, 1),
        "تغيّر_منذ_1990_pct":   round(chg_base, 1),
        "تغيّر_سنوي_pct":       round(chg_yoy, 2),
        "فجوة_NDC_Gg":         round(ndc_gap, 1),
        "وضع_NDC":             "يتجاوز الهدف" if EMISSIONS[LATEST_YEAR] > NDC["هدف_غير_مشروط_tCO2eq"]/1000 else "ضمن الهدف",
    }

# =========================================================
# رسوم لوحة التحكم
# =========================================================
def fig_bytes(fig):
    buf=io.BytesIO()
    fig.savefig(buf,format='png',dpi=150,bbox_inches='tight',facecolor='white')
    buf.seek(0); plt.close(fig); return buf

def chart_dashboard():
    """لوحة رسوم متعددة 2×2"""
    stats = compute_all()
    fig, axes = plt.subplots(2, 2, figsize=(14, 9))
    fig.suptitle(ar(f"لوحة تحكم سياسات المناخ - {PARTY['الاسم_عربي']}"),
                 fontsize=15, fontweight='bold', y=0.98)

    # رسم 1: شريط NDC Navigator
    ax1 = axes[0,0]
    routes = list(NDC_NAVIGATOR_SCORES.keys())
    scores = list(NDC_NAVIGATOR_SCORES.values())
    colors = ['#37864B' if s>=4 else '#BF8F00' if s==3 else '#C00000' for s in scores]
    ax1.bar([ar(r) for r in routes], scores, color=colors, alpha=0.85, width=0.6)
    ax1.axhline(3, color='#BF8F00', linewidth=1.5, linestyle='--', alpha=0.7)
    ax1.set_ylim(0,6); ax1.set_title(ar("تقييم NDC Navigator"), fontsize=11, fontweight='bold')
    ax1.grid(axis='y', alpha=0.3); ax1.spines[['top','right']].set_visible(False)

    # رسم 2: فجوة التمويل
    ax2 = axes[0,1]
    need = NDC["احتياج_تمويل_مليار"]
    mob  = stats["تمويل_مُعبَّأ_مليار"]
    ax2.barh([ar("الاحتياج الكلي")],[need],color='#E8F0F7',height=0.4,edgecolor='#1F497D',linewidth=1.5)
    ax2.barh([ar("الاحتياج الكلي")],[mob],color='#2E75B6',height=0.4,alpha=0.9)
    ax2.text(mob+0.5,0,ar(f"مُعبَّأ: {mob:.2f}B"),va='center',fontsize=10,color='#2E75B6',fontweight='bold')
    ax2.set_xlim(0,need*1.1); ax2.set_title(ar("الفجوة التمويلية (مليار $)"),fontsize=11,fontweight='bold')
    ax2.spines[['top','right','left']].set_visible(False); ax2.yaxis.set_visible(False)
    ax2.xaxis.set_major_formatter(matplotlib.ticker.FuncFormatter(lambda x,_:f'{x:.0f}B'))

    # رسم 3: دائري القطاعات
    ax3 = axes[1,0]
    labels = [ar(k) for k in SECTOR_SHARES_LATEST]
    values = [v*100 for v in SECTOR_SHARES_LATEST.values()]
    colors3 = ['#C00000','#2E75B6','#BF8F00','#37864B']
    wedges,_,autotexts = ax3.pie(values,labels=labels,autopct='%1.0f%%',
        colors=colors3,startangle=140,pctdistance=0.82)
    for at in autotexts: at.set_fontsize(9); at.set_fontweight('bold')
    ax3.set_title(ar(f"توزيع الانبعاثات {LATEST_YEAR}"),fontsize=11,fontweight='bold')

    # رسم 4: مقاييس سريعة
    ax4 = axes[1,1]
    ax4.axis('off')
    metrics = [
        (ar(f"إجمالي الانبعاثات {LATEST_YEAR}"), ar(f"{EMISSIONS[LATEST_YEAR]:,} Gg"), '#1F497D'),
        (ar("التغيّر منذ 1990"),       ar(f"+{stats['تغيّر_منذ_1990_pct']:.1f}%"),   '#C00000'),
        (ar("التغيّر السنوي"),          ar(f"+{stats['تغيّر_سنوي_pct']:.2f}%"),       '#BF8F00'),
        (ar("وضع NDC"),                ar(stats["وضع_NDC"]),                           '#C00000'),
        (ar("متوسط Navigator"),        ar(f"{stats['متوسط_NDC_Navigator']:.1f}/5.0"), '#BF8F00'),
        (ar("تغطية التمويل"),          ar(f"{stats['تغطية_pct']:.3f}%"),             '#C00000'),
        (ar("صفقات ITMOs نشطة"),       ar(str(ARTICLE6["صفقات_نشطة"])),               '#2E75B6'),
    ]
    for i,(label,val,color) in enumerate(metrics):
        y = 0.92 - i*0.14
        ax4.text(0.0,y,label+": ",transform=ax4.transAxes,fontsize=10,ha='left',va='top',color='#555555')
        ax4.text(1.0,y,val,transform=ax4.transAxes,fontsize=11,ha='right',va='top',
                 fontweight='bold',color=color)
    ax4.set_title(ar("المؤشرات الرئيسية"),fontsize=11,fontweight='bold')

    fig.tight_layout(rect=[0,0,1,0.96])
    return fig_bytes(fig)

# =========================================================
# دوال Word مساعدة
# =========================================================
def set_rtl(p):
    p._p.get_or_add_pPr().insert(0,OxmlElement('w:bidi'))
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT

def run(p,t,size=13,bold=False,color=None):
    r=p.add_run(t); r.font.name=FONT; r.font.size=Pt(size)
    r.font.bold=bold; r.font.color.rgb=color or C_TEXT; return r

def para(doc,t="",size=13,bold=False,color=None,sb=0,sa=5):
    p=doc.add_paragraph(); set_rtl(p)
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if t: run(p,t,size,bold,color)
    return p

def heading(doc,t,level=1):
    c={1:C_DARK,2:C_MED}; s={1:18,2:15}
    p=para(doc,t,s[level],True,c[level],sb=10,sa=4)
    if level==1:
        pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:color'),'1F497D'); pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
    return p

def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:fill'),hx); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def ctxt(cell,t,size=12,bold=False,color=None,
         align=WD_ALIGN_PARAGRAPH.RIGHT,bg=None):
    cell.text=""; p=cell.paragraphs[0]; set_rtl(p); p.alignment=align
    run(p,t,size,bold,color or C_TEXT)
    if bg: shade(cell,bg)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

def insert_chart(doc,buf,caption,width=16):
    doc.add_picture(buf,width=Cm(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=para(doc,caption,11,False,C_GRAY,sb=2,sa=8)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

# =========================================================
# بناء وثيقة لوحة التحكم الشاملة
# =========================================================
def build_dashboard():
    doc=Document(); today=datetime.date.today().strftime("%d/%m/%Y")
    stats=compute_all()

    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.right_margin=Cm(2.0); sec.left_margin=Cm(2.0)
    doc.styles['Normal'].font.name=FONT
    doc.styles['Normal'].font.size=Pt(13)

    # الغلاف
    para(doc,sb=14,sa=0)
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,'1F497D')
    p=c.paragraphs[0]; set_rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(6)
    run(p,"لوحة تحكم سياسات المناخ",24,True,C_WHITE)
    p2=c.add_paragraph(); set_rtl(p2); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after=Pt(16)
    run(p2,f"{PARTY['الاسم_عربي']} - تقرير متكامل من المنظومة الكاملة",16,False,C_WHITE)
    para(doc,f"تاريخ الإصدار: {today}  |  بيانات محدّثة: {LAST_UPDATED}  |  مؤتمر: {CURRENT_CONFERENCE['الاسم']}",
         12,False,C_GRAY,sb=8,sa=4)
    doc.add_page_break()

    # لوحة الرسوم الشاملة
    heading(doc,"لوحة المؤشرات الرئيسية",1)
    insert_chart(doc,chart_dashboard(),
        "لوحة تحكم شاملة: NDC Navigator، الفجوة التمويلية، توزيع الانبعاثات، والمؤشرات الرئيسية")
    doc.add_page_break()

    # المؤشرات الرقمية
    heading(doc,"أولاً: المؤشرات الرقمية المجمّعة",1)

    sections_data = [
        ("الجرد الوطني",[
            ("إجمالي الانبعاثات 2021",    f"{EMISSIONS[LATEST_YEAR]:,} Gg CO2eq",  "1F497D"),
            ("التغيّر منذ 1990",          f"+{stats['تغيّر_منذ_1990_pct']:.1f}%",  "C00000"),
            ("التغيّر السنوي 2020-2021",  f"+{stats['تغيّر_سنوي_pct']:.2f}%",      "BF8F00"),
        ]),
        ("المساهمة الوطنية NDC",[
            ("الهدف غير المشروط 2030", f"{NDC['هدف_غير_مشروط_tCO2eq']//1000:,} Gg", "37864B"),
            ("فجوة الانبعاثات",         f"+{stats['فجوة_NDC_Gg']:,.0f} Gg",            "C00000"),
            ("وضع الامتثال",            stats["وضع_NDC"],                               "C00000"),
        ]),
        ("التمويل المناخي",[
            ("الاحتياج الكلي",    f"{NDC['احتياج_تمويل_مليار']} مليار $",              "1F497D"),
            ("التمويل المُعبَّأ", f"{stats['تمويل_مُعبَّأ_مليار']:.3f} مليار $",       "2E75B6"),
            ("الفجوة التمويلية",  f"{stats['فجوة_تمويل_مليار']:.2f} مليار $",          "C00000"),
        ]),
        ("المادة السادسة",[
            ("صفقات ITMOs نشطة",  str(ARTICLE6["صفقات_نشطة"]),                         "2E75B6"),
            ("ITMOs سنوية",       f"{ARTICLE6['ITMOs_سنوية']:,} tCO2eq",              "1F497D"),
            ("إيرادات سنوية",     f"${ARTICLE6['إيرادات_سنوية']:,}",                   "37864B"),
        ]),
    ]

    for section_name, kpis in sections_data:
        heading(doc, section_name, 2)
        tbl=doc.add_table(2,3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        for i,(l,v,bg) in enumerate(kpis):
            ctxt(tbl.cell(0,i),l,11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,bg)
            ctxt(tbl.cell(1,i),v,15,True,RGBColor(*bytes.fromhex(bg)),WD_ALIGN_PARAGRAPH.CENTER)
        para(doc,sb=6,sa=4)

    doc.add_page_break()

    # جدول الربط بين الأدوات
    heading(doc,"ثانياً: خريطة الربط بين الأدوات",1)
    para(doc,"يوضح الجدول التالي كيف تتدفق البيانات والمخرجات بين الأدوات التسع:",13,sa=8)

    links=[
        ("shared_data.py","BTR Generator",     "انبعاثات + NDC + هوية الدولة"),
        ("shared_data.py","Article 6 Tracker", "هدف NDC للتحقق من الامتثال"),
        ("shared_data.py","GHG Analyzer",      "سنوات الجرد وبيانات الانبعاثات"),
        ("shared_data.py","Finance Tracker",   "احتياجات تمويل NDC"),
        ("shared_data.py","NDC Navigator",     "درجات تقييم المسارات السبعة"),
        ("shared_data.py","PPTX Generator",    "هوية الدولة ومجموعة التفاوض"),
        ("shared_data.py","Conference Manager","بيانات المؤتمر الحالي"),
        ("GHG Analyzer",  "BTR Generator",     "إحصاءات الجرد التفصيلية"),
        ("NDC Navigator", "PPTX Generator",    "تقييم المسارات → شرائح المواقف"),
        ("Finance Tracker","GCF Prep",         "فجوات التمويل → تحديد أولويات المشاريع"),
        ("Article 6",     "BTR Generator",     "بيانات ITMOs → قسم المادة السادسة"),
        ("NAP Central",   "NDC Navigator",     "خطط التكيف → تغذية مسار R2"),
    ]

    hdrs=["المصدر","الوجهة","البيانات المنقولة"]
    wids=[4.5,4.5,8.0]
    tbl_l=doc.add_table(1+len(links),3); tbl_l.style="Table Grid"
    tbl_l.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(h,w) in enumerate(zip(hdrs,wids)):
        cc=tbl_l.cell(0,i); cc.width=Cm(w)
        ctxt(cc,h,12,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"1F497D")
    for ri,(src,dst,data) in enumerate(links,1):
        bg="F2F7FD" if ri%2==0 else "FFFFFF"
        ctxt(tbl_l.cell(ri,0),src,11,True,C_DARK,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl_l.cell(ri,1),dst,11,True,C_MED,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl_l.cell(ri,2),data,11,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)

    doc.add_page_break()

    # حالة الأدوات
    heading(doc,"ثالثاً: حالة جميع الأدوات",1)
    tools=[
        ("BTR Generator",       "tools/btr_generator/btr_generator.py",         "✓ نشط","37864B"),
        ("Article 6 Tracker",   "tools/article6_tracker/a6_tracker.py",          "✓ نشط","37864B"),
        ("GHG Analyzer",        "tools/ghg_analyzer/ghg_analyzer.py",            "✓ نشط","37864B"),
        ("PPTX Generator",      "tools/pptx_generator/pptx_generator.py",        "✓ نشط","37864B"),
        ("Finance Tracker",     "tools/finance_tracker/finance_tracker.py",      "✓ نشط","37864B"),
        ("NDC Navigator",       "tools/ndc_navigator/ndc_navigator.py",          "✓ نشط","37864B"),
        ("NAP Central",         "tools/nap_central/nap_central.py",              "✓ نشط","37864B"),
        ("GCF Prep",            "tools/gcf_prep/gcf_prep.py",                    "✓ نشط","37864B"),
        ("Conference Manager",  "tools/conference_manager/conference_manager.py","✓ نشط","37864B"),
    ]
    hdrs2=["الأداة","المسار","الحالة"]
    wids2=[4.0,10.0,2.5]
    tbl_t=doc.add_table(1+len(tools),3); tbl_t.style="Table Grid"
    tbl_t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(h,w) in enumerate(zip(hdrs2,wids2)):
        cc=tbl_t.cell(0,i); cc.width=Cm(w)
        ctxt(cc,h,12,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"2E75B6")
    for ri,(name,path,status,sbg) in enumerate(tools,1):
        bg="F2F7FD" if ri%2==0 else "FFFFFF"
        ctxt(tbl_t.cell(ri,0),name,12,True,C_DARK,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl_t.cell(ri,1),path,10,False,C_GRAY,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl_t.cell(ri,2),status,12,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,sbg)

    # التذييل
    footer=doc.sections[-1].footer.paragraphs[0]
    set_rtl(footer); footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run(footer,f"لوحة تحكم المنظومة - {PARTY['الاسم_عربي']}  |  {today}",10,False,C_GRAY)

    return doc, stats

# =========================================================
# تحديث ملفات بيانات الأدوات من المصدر المشترك
# =========================================================
def sync_shared_data():
    """يكتب قيم مشتركة في ملفات بيانات الأدوات"""
    updated = []

    # مزامنة BTR
    btr_path = os.path.join(ROOT, "tools/btr_generator/btr_data.py")
    if os.path.exists(btr_path):
        content = open(btr_path, encoding='utf-8').read()
        # تحديث سنة التقديم إذا كانت مختلفة
        import re
        new = re.sub(r'"سنة_التقديم":\s*\d+',
                     f'"سنة_التقديم":      {datetime.date.today().year}', content)
        if new != content:
            open(btr_path,'w',encoding='utf-8').write(new)
            updated.append("btr_data.py")

    return updated

# =========================================================
# التشغيل
# =========================================================
if __name__ == "__main__":
    print("="*60)
    print("محرك ربط المنظومة - لوحة التحكم الشاملة")
    print("="*60)

    print("\n▶  مزامنة البيانات المشتركة...")
    updated = sync_shared_data()
    if updated:
        print(f"   تم تحديث: {', '.join(updated)}")
    else:
        print("   البيانات محدّثة بالفعل")

    print("\n▶  توليد لوحة التحكم الشاملة...")
    doc, stats = build_dashboard()
    out_doc = "pipeline/output_DASHBOARD_لوحة_التحكم.docx"
    os.makedirs("pipeline", exist_ok=True)
    doc.save(out_doc)
    print(f"   ✓ تم إنشاء لوحة التحكم: {out_doc}")

    # حفظ الحالة الكاملة JSON
    state = {
        "تاريخ_الإنشاء":        datetime.datetime.now().isoformat(),
        "الدولة":                PARTY["الاسم_عربي"],
        "المؤتمر_الحالي":       CURRENT_CONFERENCE["الاسم"],
        "المؤشرات_المجمّعة":    stats,
        "الانبعاثات":            EMISSIONS,
        "NDC":                   {k:v for k,v in NDC.items()},
        "التمويل":               FINANCE_MOBILIZED,
        "المادة_السادسة":        ARTICLE6,
        "تقييم_Navigator":       NDC_NAVIGATOR_SCORES,
        "الأدوات_النشطة":        9,
    }
    out_json = "pipeline/output_STATE_الحالة_الكاملة.json"
    with open(out_json,'w',encoding='utf-8') as f:
        json.dump(state,f,ensure_ascii=False,indent=2)
    print(f"   ✓ تم حفظ الحالة: {out_json}")

    print("\n" + "="*60)
    print("ملخص الحالة")
    print("="*60)
    print(f"  إجمالي الانبعاثات {LATEST_YEAR}: {EMISSIONS[LATEST_YEAR]:,} Gg CO2eq")
    print(f"  وضع NDC:               {stats['وضع_NDC']}")
    print(f"  فجوة التمويل:          {stats['فجوة_تمويل_مليار']:.2f} مليار $")
    print(f"  تغطية التمويل:         {stats['تغطية_pct']:.3f}%")
    print(f"  متوسط Navigator:       {stats['متوسط_NDC_Navigator']:.1f}/5.0")
    print(f"  صفقات ITMOs نشطة:      {ARTICLE6['صفقات_نشطة']}")
    print(f"  المؤتمر القادم:         {CURRENT_CONFERENCE['الاسم']}")
