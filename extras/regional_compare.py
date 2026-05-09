"""
أداة المقارنة الإقليمية التلقائية
التشغيل: python extras/regional_compare.py
تولّد تقرير مقارنة شهري بين دول المنطقة
"""

import sys, os, io, json, datetime
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, 'core'))
from shared_data import PARTY, EMISSIONS, LATEST_YEAR, NDC

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import arabic_reshaper
from bidi.algorithm import get_display
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

plt.rcParams['font.family'] = 'DejaVu Sans'
def ar(t): return get_display(arabic_reshaper.reshape(str(t)))

C_DARK=RGBColor(0x1F,0x49,0x7D); C_MED=RGBColor(0x2E,0x75,0xB6)
C_WHITE=RGBColor(0xFF,0xFF,0xFF); C_GRAY=RGBColor(0x70,0x70,0x70)
C_TEXT=RGBColor(0x26,0x26,0x26); FONT="Simplified Arabic"

# ===== بيانات الدول الإقليمية =====
REGIONAL_DATA = {
    "العراق":                    {"انبعاثات_2020": 190342, "NDC_pct": 2.5,  "تمويل_مليار": 88.0,  "per_capita": 4.6,  "GDP_intensity": 0.82},
    "المملكة العربية السعودية":  {"انبعاثات_2020": 700100, "NDC_pct": 30.0, "تمويل_مليار": 0.0,   "per_capita": 19.8, "GDP_intensity": 0.45},
    "إيران":                     {"انبعاثات_2020": 749000, "NDC_pct": 4.0,  "تمويل_مليار": 17.5,  "per_capita": 8.8,  "GDP_intensity": 0.95},
    "مصر":                       {"انبعاثات_2020": 340200, "NDC_pct": 33.0, "تمويل_مليار": 196.0, "per_capita": 3.2,  "GDP_intensity": 0.38},
    "الجزائر":                   {"انبعاثات_2020": 194000, "NDC_pct": 22.0, "تمويل_مليار": 7.0,   "per_capita": 4.4,  "GDP_intensity": 0.52},
    "الإمارات":                  {"انبعاثات_2020": 256000, "NDC_pct": 19.0, "تمويل_مليار": 0.0,   "per_capita": 25.8, "GDP_intensity": 0.35},
    "الأردن":                    {"انبعاثات_2020": 28800,  "NDC_pct": 14.0, "تمويل_مليار": 5.4,   "per_capita": 2.7,  "GDP_intensity": 0.28},
    "الكويت":                    {"انبعاثات_2020": 98000,  "NDC_pct": 7.4,  "تمويل_مليار": 0.0,   "per_capita": 22.4, "GDP_intensity": 0.68},
}

def fig_bytes(fig):
    buf=io.BytesIO()
    fig.savefig(buf,format='png',dpi=150,bbox_inches='tight',facecolor='white')
    buf.seek(0); plt.close(fig); return buf

def chart_emissions_compare():
    countries = list(REGIONAL_DATA.keys())
    values    = [REGIONAL_DATA[c]["انبعاثات_2020"] for c in countries]
    colors    = ['#C00000' if c=="العراق" else '#2E75B6' for c in countries]

    sorted_pairs = sorted(zip(values, countries, colors))
    values, countries, colors = zip(*sorted_pairs)

    fig, ax = plt.subplots(figsize=(10,6))
    bars = ax.barh([ar(c) for c in countries], values, color=colors, alpha=0.85, height=0.6)
    for bar, val in zip(bars, values):
        ax.text(bar.get_width()+2000, bar.get_y()+bar.get_height()/2,
                f'{val/1000:.0f}k', va='center', fontsize=9)
    ax.set_xlabel(ar("Gg CO2eq"), fontsize=11)
    ax.set_title(ar("مقارنة الانبعاثات الإقليمية - 2020"), fontsize=13, fontweight='bold')
    ax.xaxis.set_major_formatter(
        matplotlib.ticker.FuncFormatter(lambda x,_:f'{x/1000:.0f}k'))
    ax.grid(axis='x', alpha=0.3)
    ax.spines[['top','right']].set_visible(False)
    fig.tight_layout()
    return fig_bytes(fig)

def chart_ndc_ambition():
    countries = list(REGIONAL_DATA.keys())
    ndcs      = [REGIONAL_DATA[c]["NDC_pct"] for c in countries]
    colors    = ['#C00000' if c=="العراق" else '#70AD47' for c in countries]

    sorted_pairs = sorted(zip(ndcs, countries, colors))
    ndcs, countries, colors = zip(*sorted_pairs)

    fig, ax = plt.subplots(figsize=(10,6))
    bars = ax.barh([ar(c) for c in countries], ndcs, color=colors, alpha=0.85, height=0.6)
    ax.axvline(15, color='#BF8F00', linewidth=1.5, linestyle='--', alpha=0.8,
               label=ar("متوسط الدول النامية (15%)"))
    for bar, val in zip(bars, ndcs):
        ax.text(bar.get_width()+0.3, bar.get_y()+bar.get_height()/2,
                f'{val}%', va='center', fontsize=9)
    ax.set_xlabel(ar("هدف الخفض %"), fontsize=11)
    ax.set_title(ar("مقارنة طموح NDC الإقليمية"), fontsize=13, fontweight='bold')
    ax.legend(fontsize=10)
    ax.grid(axis='x', alpha=0.3)
    ax.spines[['top','right']].set_visible(False)
    fig.tight_layout()
    return fig_bytes(fig)

def chart_per_capita():
    countries  = list(REGIONAL_DATA.keys())
    per_capita = [REGIONAL_DATA[c]["per_capita"] for c in countries]
    colors     = ['#C00000' if c=="العراق" else '#4472C4' for c in countries]

    fig, ax = plt.subplots(figsize=(10,5))
    ax.bar([ar(c)[:8] for c in countries], per_capita, color=colors, alpha=0.85, width=0.6)
    for i, (val, c) in enumerate(zip(per_capita, countries)):
        ax.text(i, val+0.1, f'{val}', ha='center', fontsize=9, fontweight='bold')
    ax.set_ylabel(ar("طن CO2eq / فرد"), fontsize=11)
    ax.set_title(ar("انبعاثات الفرد مقارنةً إقليمياً"), fontsize=13, fontweight='bold')
    ax.grid(axis='y', alpha=0.3)
    ax.spines[['top','right']].set_visible(False)
    plt.xticks(rotation=25, ha='right', fontsize=9)
    fig.tight_layout()
    return fig_bytes(fig)

# ===== دوال Word =====
def set_rtl(p):
    p._p.get_or_add_pPr().insert(0,OxmlElement('w:bidi'))
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
def run(p,t,size=13,bold=False,color=None):
    r=p.add_run(t); r.font.name=FONT; r.font.size=Pt(size)
    r.font.bold=bold; r.font.color.rgb=color or C_TEXT; return r
def para(doc,t="",size=13,bold=False,color=None,sb=0,sa=5):
    p=doc.add_paragraph(); set_rtl(p)
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if t: run(p,t,size,bold,color)
    return p
def heading(doc,t,level=1):
    c={1:C_DARK,2:C_MED}; s={1:18,2:15}
    p=para(doc,t,s[level],True,c[level],sb=10,sa=4)
    if level==1:
        pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:color'),'1F497D'); pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
    return p
def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:fill'),hx); shd.set(qn('w:val'),'clear'); tcPr.append(shd)
def ctxt(cell,t,size=12,bold=False,color=None,align=WD_ALIGN_PARAGRAPH.RIGHT,bg=None):
    cell.text=""; p=cell.paragraphs[0]; set_rtl(p); p.alignment=align
    run(p,t,size,bold,color or C_TEXT)
    if bg: shade(cell,bg)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
def insert_chart(doc,buf,caption,width=15):
    doc.add_picture(buf,width=Cm(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=para(doc,caption,11,False,C_GRAY,sb=2,sa=8)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

def build_report():
    doc=Document(); today=datetime.date.today().strftime("%d/%m/%Y")

    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.right_margin=Cm(2.0); sec.left_margin=Cm(2.0)
    doc.styles['Normal'].font.name=FONT
    doc.styles['Normal'].font.size=Pt(13)

    # الغلاف
    para(doc,sb=18,sa=0)
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,'1F497D')
    p=c.paragraphs[0]; set_rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(6)
    run(p,"تقرير المقارنة الإقليمية لسياسات المناخ",22,True,C_WHITE)
    p2=c.add_paragraph(); set_rtl(p2); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after=Pt(16)
    run(p2,f"العراق ودول المنطقة - {today}",15,False,C_WHITE)
    doc.add_page_break()

    # جدول المقارنة
    heading(doc,"أولاً: جدول المقارنة الشاملة",1)
    countries = list(REGIONAL_DATA.keys())
    hdrs=["الدولة","الانبعاثات 2020\n(ألف Gg)","هدف NDC %","طموح التمويل\n(مليار $)","نصيب الفرد\n(طن/فرد)"]
    wids=[4.5,3.5,3.0,3.5,3.0]
    tbl=doc.add_table(1+len(countries),5); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(h,w) in enumerate(zip(hdrs,wids)):
        cc=tbl.cell(0,i); cc.width=Cm(w)
        ctxt(cc,h,11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"1F497D")
    for ri,country in enumerate(countries,1):
        d=REGIONAL_DATA[country]
        bg="FFE7E7" if country=="العراق" else ("F2F7FD" if ri%2==0 else "FFFFFF")
        bold_flag = country=="العراق"
        ctxt(tbl.cell(ri,0),country,12,bold_flag,C_DARK if bold_flag else C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl.cell(ri,1),f"{d['انبعاثات_2020']/1000:.1f}",12,bold_flag,C_DARK if bold_flag else C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl.cell(ri,2),f"{d['NDC_pct']}%",12,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl.cell(ri,3),f"{d['تمويل_مليار']}" if d["تمويل_مليار"]>0 else "—",12,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl.cell(ri,4),f"{d['per_capita']}",12,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)

    doc.add_page_break()

    # الرسوم البيانية
    heading(doc,"ثانياً: الرسوم البيانية المقارنة",1)
    insert_chart(doc,chart_emissions_compare(),"شكل (1): مقارنة الانبعاثات الإقليمية 2020")
    insert_chart(doc,chart_ndc_ambition(),"شكل (2): مقارنة طموح NDC الإقليمية")
    insert_chart(doc,chart_per_capita(),"شكل (3): انبعاثات الفرد مقارنةً إقليمياً",width=14)

    # موقع العراق
    doc.add_page_break()
    heading(doc,"ثالثاً: الموقع الإقليمي للعراق",1)
    iraq = REGIONAL_DATA["العراق"]
    all_em = sorted(REGIONAL_DATA.values(),key=lambda x:x["انبعاثات_2020"],reverse=True)
    rank_em = next(i+1 for i,d in enumerate(all_em) if d["انبعاثات_2020"]==iraq["انبعاثات_2020"])
    all_ndc = sorted(REGIONAL_DATA.values(),key=lambda x:x["NDC_pct"],reverse=True)
    rank_ndc = next(i+1 for i,d in enumerate(all_ndc) if d["NDC_pct"]==iraq["NDC_pct"])

    findings=[
        f"يحتل العراق المرتبة {rank_em} إقليمياً من حيث إجمالي الانبعاثات",
        f"هدف NDC العراقي ({iraq['NDC_pct']}%) يُصنَّف ضمن الأدنى إقليمياً - المرتبة {rank_ndc}",
        f"نصيب الفرد ({iraq['per_capita']} طن) أدنى من السعودية والإمارات والكويت",
        "العراق يتصدر احتياجات التمويل (88 مليار $) مما يعكس حجم التحدي التنموي",
        "رفع هدف NDC إلى 10-15% سيُحسّن الموقع التفاوضي إقليمياً ودولياً",
    ]
    for i,f in enumerate(findings,1):
        p=doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after=Pt(5)
        run(p,f"{i}.  ",13,True,C_MED); run(p,f,13,False,C_TEXT)

    footer=doc.sections[-1].footer.paragraphs[0]
    set_rtl(footer); footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run(footer,f"تقرير المقارنة الإقليمية - العراق  |  {today}",10,False,C_GRAY)
    return doc

if __name__ == "__main__":
    print("="*55)
    print("  أداة المقارنة الإقليمية التلقائية")
    print("="*55)
    print("  جارٍ توليد تقرير المقارنة...")
    doc=build_report()
    out=os.path.join(ROOT,"extras","output_REGIONAL_المقارنة_الإقليمية.docx")
    doc.save(out)
    print(f"  ✓  تم إنشاء التقرير: {out}")
