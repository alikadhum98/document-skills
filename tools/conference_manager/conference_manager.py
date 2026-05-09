"""
أداة إدارة وثائق المشاركة في مؤتمرات المناخ
التشغيل: python tools/conference_manager/conference_manager.py
المخرجات: ورقة إحاطة Word + جدول الجلسات + متابعة الإجراءات JSON
"""

import sys, os, json, datetime
sys.path.insert(0, os.path.dirname(__file__))
from conf_data import CONFERENCE, SESSIONS, ACTION_ITEMS

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_DARK=RGBColor(0x1F,0x49,0x7D); C_MED=RGBColor(0x2E,0x75,0xB6)
C_GREEN=RGBColor(0x37,0x86,0x44); C_RED=RGBColor(0xC0,0x00,0x00)
C_AMBER=RGBColor(0xBF,0x8F,0x00); C_WHITE=RGBColor(0xFF,0xFF,0xFF)
C_GRAY=RGBColor(0x70,0x70,0x70); C_TEXT=RGBColor(0x26,0x26,0x26)
FONT="Simplified Arabic"

PRIORITY_BG={"عالية":"C00000","متوسطة":"BF8F00","منخفضة":"37864B"}
STATUS_BG={"مجدولة":"2E75B6","منتهية":"37864B","ملغاة":"C00000","معلّق":"BF8F00","مكتمل":"37864B"}
BODY_COLORS={"SBSTA":"1F497D","SBI":"2E75B6","Arab Group":"37864B","CMA":"7030A0","COP":"C00000"}

def set_rtl(p):
    p._p.get_or_add_pPr().insert(0,OxmlElement('w:bidi'))
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT

def run(p,t,size=13,bold=False,color=None):
    r=p.add_run(t); r.font.name=FONT; r.font.size=Pt(size)
    r.font.bold=bold; r.font.color.rgb=color or C_TEXT; return r

def para(doc,t="",size=13,bold=False,color=None,sb=0,sa=5):
    p=doc.add_paragraph(); set_rtl(p)
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if t: run(p,t,size,bold,color)
    return p

def heading(doc,t,level=1):
    c={1:C_DARK,2:C_MED,3:C_TEXT}; s={1:18,2:15,3:13}
    p=para(doc,t,s[level],True,c[level],sb=10,sa=4)
    if level==1:
        pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:color'),'1F497D'); pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
    return p

def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:fill'),hx); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def ctxt(cell,t,size=12,bold=False,color=None,
         align=WD_ALIGN_PARAGRAPH.RIGHT,bg=None):
    cell.text=""; p=cell.paragraphs[0]; set_rtl(p); p.alignment=align
    run(p,t,size,bold,color or C_TEXT)
    if bg: shade(cell,bg)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

def build_briefing():
    doc=Document(); today=datetime.date.today().strftime("%d/%m/%Y")
    conf=CONFERENCE

    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.right_margin=Cm(2.0); sec.left_margin=Cm(2.0)
    doc.styles['Normal'].font.name=FONT
    doc.styles['Normal'].font.size=Pt(13)

    # الغلاف
    para(doc,sb=16,sa=0)
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,'1F497D')
    p=c.paragraphs[0]; set_rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
    run(p,"ورقة إحاطة الوفد العراقي",20,True,C_WHITE)
    p2=c.add_paragraph(); set_rtl(p2); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after=Pt(14)
    run(p2,conf["الاسم"],16,True,C_WHITE)

    para(doc,f"المكان: {conf['المكان']}  |  التاريخ: {conf['التاريخ']}  |  إعداد: {today}",
         12,False,C_GRAY,sb=8,sa=4)
    doc.add_page_break()

    # معلومات المؤتمر
    heading(doc,"أولاً: معلومات المؤتمر",1)
    rows=[
        ("اسم المؤتمر",   conf["الاسم"]),
        ("المكان",         conf["المكان"]),
        ("الفترة",         f"{conf['التاريخ_من']} إلى {conf['التاريخ_إلى']}"),
        ("حجم الوفد",      f"{len(conf['الوفد_العراقي'])} أعضاء"),
    ]
    tbl=doc.add_table(len(rows),2); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ri,(l,v) in enumerate(rows):
        bg="EEF4FB" if ri%2==0 else "FFFFFF"
        ctxt(tbl.cell(ri,0),l,12,True,C_DARK,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl.cell(ri,1),v,12,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)
    para(doc,sb=6,sa=4)

    # تشكيل الوفد
    heading(doc,"ثانياً: تشكيل الوفد العراقي",1)
    hdrs=["الاسم","المنصب","الوزارة"]
    wids=[5,6,6]
    tbl2=doc.add_table(1+len(conf["الوفد_العراقي"]),3); tbl2.style="Table Grid"
    tbl2.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(h,w) in enumerate(zip(hdrs,wids)):
        cc=tbl2.cell(0,i); cc.width=Cm(w)
        ctxt(cc,h,12,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"1F497D")
    for ri,member in enumerate(conf["الوفد_العراقي"],1):
        bg="F2F7FD" if ri%2==0 else "FFFFFF"
        ctxt(tbl2.cell(ri,0),member["الاسم"],12,True,C_DARK,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl2.cell(ri,1),member["المنصب"],12,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl2.cell(ri,2),member["الوزارة"],12,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
    para(doc,sb=6,sa=4)

    # جدول الجلسات
    heading(doc,"ثالثاً: جدول الجلسات والأولويات",1)
    hdrs3=["المعرّف","الموضوع","الهيئة","التاريخ","الوقت","الأولوية"]
    wids3=[2.5,6.0,2.8,2.5,1.8,2.2]
    tbl3=doc.add_table(1+len(SESSIONS),6); tbl3.style="Table Grid"
    tbl3.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(h,w) in enumerate(zip(hdrs3,wids3)):
        cc=tbl3.cell(0,i); cc.width=Cm(w)
        ctxt(cc,h,11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"2E75B6")
    for ri,s in enumerate(SESSIONS,1):
        bg="F2F7FD" if ri%2==0 else "FFFFFF"
        pbg=PRIORITY_BG.get(s["الأولوية"],"888888")
        bbg=BODY_COLORS.get(s["الهيئة"],"888888")
        ctxt(tbl3.cell(ri,0),s["المعرّف"],11,True,C_DARK,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl3.cell(ri,1),s["الموضوع"],11,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl3.cell(ri,2),s["الهيئة"],11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,bbg)
        ctxt(tbl3.cell(ri,3),s["التاريخ"],11,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl3.cell(ri,4),s["الوقت"],11,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl3.cell(ri,5),s["الأولوية"],11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,pbg)

    doc.add_page_break()

    # المواقف التفاوضية
    heading(doc,"رابعاً: المواقف التفاوضية لكل جلسة",1)
    for s in [x for x in SESSIONS if x["الأولوية"]=="عالية"]:
        bbg=BODY_COLORS.get(s["الهيئة"],"888888")
        heading(doc,f"{s['المعرّف']}: {s['الموضوع'][:45]}",2)
        p=doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after=Pt(4)
        run(p,"موقف العراق: ",13,True,C_MED)
        run(p,s["موقف_العراق"],13,False,C_TEXT)
        para(doc,"النقاط الرئيسية للتدخل:",13,True,C_DARK,sb=4,sa=2)
        for pt in s["النقاط_الرئيسية"]:
            p=doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after=Pt(3)
            run(p,"◀  "+pt,12,False,C_TEXT)
        if s["الوثائق"]:
            p=doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after=Pt(8)
            run(p,"الوثائق: ",12,True,C_GRAY)
            run(p,"  |  ".join(s["الوثائق"]),12,False,C_GRAY)

    # نقاط الإجراءات
    heading(doc,"خامساً: نقاط الإجراءات والمتابعة",1)
    hdrs4=["المعرّف","الإجراء المطلوب","المسؤول","الموعد","الأولوية","الحالة"]
    wids4=[2.0,6.5,2.5,2.5,2.0,2.0]
    tbl4=doc.add_table(1+len(ACTION_ITEMS),6); tbl4.style="Table Grid"
    tbl4.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(h,w) in enumerate(zip(hdrs4,wids4)):
        cc=tbl4.cell(0,i); cc.width=Cm(w)
        ctxt(cc,h,11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"1F497D")
    for ri,act in enumerate(ACTION_ITEMS,1):
        bg="F2F7FD" if ri%2==0 else "FFFFFF"
        pbg=PRIORITY_BG.get(act["الأولوية"],"888888")
        sbg=STATUS_BG.get(act["الحالة"],"888888")
        ctxt(tbl4.cell(ri,0),act["المعرّف"],11,True,C_DARK,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl4.cell(ri,1),act["الإجراء"],11,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl4.cell(ri,2),act["المسؤول"],11,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl4.cell(ri,3),act["الموعد"],11,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl4.cell(ri,4),act["الأولوية"],11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,pbg)
        ctxt(tbl4.cell(ri,5),act["الحالة"],11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,sbg)

    footer=doc.sections[-1].footer.paragraphs[0]
    set_rtl(footer); footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run(footer,f"ورقة إحاطة - {conf['الاسم']}  |  {today}  |  سري للاستخدام الرسمي",
        10,False,C_GRAY)

    return doc

if __name__=="__main__":
    print("جارٍ إنشاء ورقة إحاطة المؤتمر...")
    doc=build_briefing()
    out="tools/conference_manager/output_CONF_ورقة_الإحاطة.docx"
    doc.save(out)
    print(f"✓ تم إنشاء ورقة الإحاطة: {out}")

    # حفظ JSON للمتابعة
    tracker={"المؤتمر":CONFERENCE["الاسم"],
             "الجلسات":len(SESSIONS),
             "إجراءات_معلّقة":sum(1 for a in ACTION_ITEMS if a["الحالة"]=="معلّق"),
             "الإجراءات":ACTION_ITEMS}
    out_j="tools/conference_manager/output_CONF_متابعة_الإجراءات.json"
    with open(out_j,'w',encoding='utf-8') as f:
        json.dump(tracker,f,ensure_ascii=False,indent=2)
    print(f"✓ تم حفظ متابعة الإجراءات: {out_j}")
    print(f"  الجلسات: {len(SESSIONS)}  |  إجراءات معلّقة: {tracker['إجراءات_معلّقة']}")
