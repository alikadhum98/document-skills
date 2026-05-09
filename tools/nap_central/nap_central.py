"""
أداة مولّد الخطة الوطنية للتكيف NAP - العراق
التشغيل: python tools/nap_central/nap_central.py
"""

import sys, os, io, json, datetime
sys.path.insert(0, os.path.dirname(__file__))
from nap_data import NAP_INFO, SECTORS, FINANCING, MONITORING

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import arabic_reshaper
from bidi.algorithm import get_display
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

plt.rcParams['font.family'] = 'DejaVu Sans'
def ar(t): return get_display(arabic_reshaper.reshape(str(t)))

C_DARK=RGBColor(0x1F,0x49,0x7D); C_MED=RGBColor(0x2E,0x75,0xB6)
C_GREEN=RGBColor(0x37,0x86,0x44); C_RED=RGBColor(0xC0,0x00,0x00)
C_AMBER=RGBColor(0xBF,0x8F,0x00); C_WHITE=RGBColor(0xFF,0xFF,0xFF)
C_GRAY=RGBColor(0x70,0x70,0x70); C_TEXT=RGBColor(0x26,0x26,0x26)
FONT="Simplified Arabic"

VUL_COLORS={"بالغة":"C00000","عالية":"BF8F00","متوسطة":"2E75B6","منخفضة":"37864B"}

def fig_bytes(fig):
    buf=io.BytesIO()
    fig.savefig(buf,format='png',dpi=150,bbox_inches='tight',facecolor='white')
    buf.seek(0); plt.close(fig); return buf

def chart_budget():
    labels=[ar(s["القطاع"]) for s in SECTORS]
    values=[s["التمويل_مليون"] for s in SECTORS]
    colors=['#C00000','#BF8F00','#2E75B6','#37864B']

    fig,ax=plt.subplots(figsize=(8,5))
    bars=ax.bar(labels,values,color=colors,alpha=0.85,width=0.6)
    for bar,val in zip(bars,values):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+10,
                f'${val:,}M',ha='center',fontsize=10,fontweight='bold')
    ax.set_ylabel(ar("مليون دولار"),fontsize=11)
    ax.set_title(ar("توزيع ميزانية NAP حسب القطاع"),fontsize=13,fontweight='bold')
    ax.grid(axis='y',alpha=0.3)
    ax.spines[['top','right']].set_visible(False)
    fig.tight_layout()
    return fig_bytes(fig)

def chart_financing():
    labels=[ar("محلي"),ar("GCF"),ar("GEF"),ar("ثنائي"),ar("خاص")]
    values=[FINANCING["محلي_pct"],FINANCING["GCF_pct"],FINANCING["GEF_pct"],
            FINANCING["ثنائي_pct"],FINANCING["قطاع_خاص_pct"]]
    colors=['#1F497D','#2E75B6','#4472C4','#70AD47','#BF8F00']

    fig,ax=plt.subplots(figsize=(7,6))
    wedges,texts,autotexts=ax.pie(values,labels=labels,autopct='%1.0f%%',
        colors=colors,startangle=140,pctdistance=0.82,explode=[0.04]*5)
    for at in autotexts:
        at.set_fontsize(11); at.set_fontweight('bold')
    ax.set_title(ar("مصادر تمويل NAP"),fontsize=13,fontweight='bold')
    fig.tight_layout()
    return fig_bytes(fig)

def set_rtl(p):
    p._p.get_or_add_pPr().insert(0,OxmlElement('w:bidi'))
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT

def run(p,t,size=13,bold=False,color=None):
    r=p.add_run(t); r.font.name=FONT; r.font.size=Pt(size)
    r.font.bold=bold; r.font.color.rgb=color or C_TEXT; return r

def para(doc,t="",size=13,bold=False,color=None,sb=0,sa=5):
    p=doc.add_paragraph(); set_rtl(p)
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if t: run(p,t,size,bold,color)
    return p

def heading(doc,t,level=1):
    c={1:C_DARK,2:C_MED,3:C_TEXT}; s={1:18,2:15,3:13}
    p=para(doc,t,s[level],True,c[level],sb=10,sa=4)
    if level==1:
        pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:color'),'1F497D'); pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
    return p

def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:fill'),hx); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def ctxt(cell,t,size=12,bold=False,color=None,
         align=WD_ALIGN_PARAGRAPH.RIGHT,bg=None):
    cell.text=""; p=cell.paragraphs[0]; set_rtl(p); p.alignment=align
    run(p,t,size,bold,color or C_TEXT)
    if bg: shade(cell,bg)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

def insert_chart(doc,buf,caption,width=14):
    doc.add_picture(buf,width=Cm(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=para(doc,caption,11,False,C_GRAY,sb=2,sa=8)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

def build_report():
    doc=Document(); today=datetime.date.today().strftime("%d/%m/%Y")

    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.right_margin=Cm(2.0); sec.left_margin=Cm(2.0)
    doc.styles['Normal'].font.name=FONT
    doc.styles['Normal'].font.size=Pt(13)

    # الغلاف
    para(doc,sb=20,sa=0)
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,'1F497D')
    p=c.paragraphs[0]; set_rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(6)
    run(p,NAP_INFO["العنوان"],22,True,C_WHITE)
    p2=c.add_paragraph(); set_rtl(p2); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after=Pt(16)
    run(p2,f"{NAP_INFO['الدولة']} | {NAP_INFO['الفترة']}",15,False,C_WHITE)
    para(doc,f"الجهة القائدة: {NAP_INFO['الجهة_القائدة']}  |  {today}",12,False,C_GRAY,sb=8,sa=4)
    doc.add_page_break()

    # الملخص
    heading(doc,"أولاً: الملخص التنفيذي",1)
    total=sum(s["التمويل_مليون"] for s in SECTORS)
    kpis=[
        ("عدد القطاعات المشمولة",str(len(SECTORS)),"2E75B6"),
        ("إجمالي الميزانية",f"${NAP_INFO['الميزانية_مليار']} مليار","1F497D"),
        ("الفترة الزمنية",NAP_INFO["الفترة"],"37864B"),
    ]
    tbl=doc.add_table(2,3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(l,v,bg) in enumerate(kpis):
        ctxt(tbl.cell(0,i),l,12,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl.cell(1,i),v,16,True,RGBColor(*bytes.fromhex(bg)),WD_ALIGN_PARAGRAPH.CENTER)

    para(doc,sb=8,sa=4)
    insert_chart(doc,chart_budget(),"شكل (1): توزيع ميزانية NAP حسب القطاع")
    insert_chart(doc,chart_financing(),"شكل (2): مصادر تمويل NAP",width=11)
    doc.add_page_break()

    # تفاصيل القطاعات
    heading(doc,"ثانياً: خطط التكيف القطاعية",1)
    for sec in SECTORS:
        vbg=VUL_COLORS.get(sec["درجة_الهشاشة"],"888888")
        heading(doc,f"{sec['الأولوية']}. قطاع {sec['القطاع']}",2)

        # بطاقة القطاع
        tbl2=doc.add_table(1,4); tbl2.style="Table Grid"; tbl2.alignment=WD_TABLE_ALIGNMENT.CENTER
        ctxt(tbl2.cell(0,0),"درجة الهشاشة",11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,vbg)
        ctxt(tbl2.cell(0,1),sec["درجة_الهشاشة"],13,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,vbg)
        ctxt(tbl2.cell(0,2),"الميزانية",11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"1F497D")
        ctxt(tbl2.cell(0,3),f"${sec['التمويل_مليون']:,}M",13,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"1F497D")

        # المخاطر
        para(doc,"المخاطر المناخية الرئيسية:",13,True,C_RED,sb=6,sa=2)
        for risk in sec["المخاطر"]:
            p=doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after=Pt(3)
            run(p,"⚠  "+risk,12,False,C_TEXT)

        # إجراءات التكيف
        para(doc,"إجراءات التكيف:",13,True,C_GREEN,sb=6,sa=2)
        hdrs=["الإجراء","التكلفة (م$)","الموعد","الجهة المسؤولة"]
        wids=[6.0,2.5,3.0,3.5]
        tbl3=doc.add_table(1+len(sec["الإجراءات"]),4)
        tbl3.style="Table Grid"; tbl3.alignment=WD_TABLE_ALIGNMENT.CENTER
        for i,(h,w) in enumerate(zip(hdrs,wids)):
            c=tbl3.cell(0,i); c.width=Cm(w)
            ctxt(c,h,11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"2E75B6")
        for ri,act in enumerate(sec["الإجراءات"],1):
            bg="F2F7FD" if ri%2==0 else "FFFFFF"
            ctxt(tbl3.cell(ri,0),act["الإجراء"],11,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)
            ctxt(tbl3.cell(ri,1),f"{act['التكلفة_مليون']}",12,True,C_DARK,WD_ALIGN_PARAGRAPH.CENTER,bg)
            ctxt(tbl3.cell(ri,2),act["الموعد"],11,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
            ctxt(tbl3.cell(ri,3),act["الجهة"],11,False,C_GRAY,WD_ALIGN_PARAGRAPH.RIGHT,bg)

        # المؤشر
        p=doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after=Pt(10)
        run(p,"مؤشر 2030: ",13,True,C_MED)
        run(p,sec["المؤشر_2030"],13,False,C_TEXT)

    doc.add_page_break()

    # المتابعة
    heading(doc,"ثالثاً: إطار المتابعة والتقييم",1)
    mon_rows=[
        ("دورة الإبلاغ",MONITORING["دورة_الإبلاغ"]),
        ("الربط مع BTR",  "نعم - يُدرَج في تقارير الشفافية السنوية"),
        ("الربط مع NDC",  "نعم - مؤشرات NAP تُغذّي أهداف NDC"),
        ("SDGs المرتبطة", "  |  ".join(MONITORING["SDGs"])),
    ]
    tbl4=doc.add_table(len(mon_rows),2); tbl4.style="Table Grid"; tbl4.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ri,(l,v) in enumerate(mon_rows):
        bg="EEF4FB" if ri%2==0 else "FFFFFF"
        ctxt(tbl4.cell(ri,0),l,12,True,C_DARK,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl4.cell(ri,1),v,12,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)

    footer=doc.sections[-1].footer.paragraphs[0]
    set_rtl(footer); footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run(footer,f"{NAP_INFO['العنوان']} - {NAP_INFO['الدولة']}  |  {today}",10,False,C_GRAY)

    return doc

if __name__=="__main__":
    print("جارٍ إنشاء وثيقة NAP...")
    doc=build_report()
    out="tools/nap_central/output_NAP_الخطة_الوطنية_للتكيف.docx"
    doc.save(out)
    total=sum(s["التمويل_مليون"] for s in SECTORS)
    print(f"✓ تم إنشاء الوثيقة: {out}")
    print(f"  القطاعات: {len(SECTORS)}  |  الميزانية الكلية: ${total:,}M")
