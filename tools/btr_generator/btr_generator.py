"""
مولّد تقرير الشفافية الثنائي (BTR) العربي
يقرأ البيانات من btr_data.py ويولّد وثيقة Word منسّقة رسمياً
التشغيل: python tools/btr_generator/btr_generator.py
"""

import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from btr_data import BTR_DATA

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ========== ألوان الهوية البصرية ==========
C_DARK   = RGBColor(0x1F, 0x49, 0x7D)   # أزرق داكن
C_MED    = RGBColor(0x2E, 0x75, 0xB6)   # أزرق متوسط
C_LIGHT  = RGBColor(0xDE, 0xEB, 0xF7)   # أزرق فاتح
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT   = RGBColor(0x26, 0x26, 0x26)
C_GRAY   = RGBColor(0x70, 0x70, 0x70)
C_GREEN  = RGBColor(0x37, 0x86, 0x44)
C_RED    = RGBColor(0xC0, 0x00, 0x00)
FONT     = "Simplified Arabic"

# ========== دوال مساعدة ==========
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def add_run(paragraph, text, size=14, bold=False, color=None, italic=False):
    run = paragraph.add_run(text)
    run.font.name    = FONT
    run.font.size    = Pt(size)
    run.font.bold    = bold
    run.font.italic  = italic
    run.font.color.rgb = color if color else C_TEXT
    return run

def add_para(doc, text="", size=14, bold=False, color=None,
             space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        add_run(p, text, size, bold, color, italic)
    return p

def add_heading(doc, text, level=1):
    colors = {1: C_DARK, 2: C_MED, 3: C_TEXT}
    sizes  = {1: 18,     2: 15,    3: 14}
    p = add_para(doc, text, size=sizes[level], bold=True,
                 color=colors[level], space_before=10, space_after=4)
    if level == 1:
        # خط أسفل العنوان
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def shade_cell(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)

def set_cell_text(cell, text, size=12, bold=False, color=None,
                  align=WD_ALIGN_PARAGRAPH.RIGHT, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = align
    add_run(p, text, size, bold, color or C_TEXT)
    if bg:
        shade_cell(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

# ========== بناء الوثيقة ==========
def build_btr(data):
    doc  = Document()
    info = data["معلومات_التقرير"]

    # --- إعداد الصفحة ---
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.left_margin   = Cm(2.5)

    # --- الخط الافتراضي ---
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(14)

    # =========================================
    # صفحة الغلاف
    # =========================================
    add_para(doc, space_before=30, space_after=0)
    add_para(doc, space_before=20, space_after=0)

    # شريط علوي
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade_cell(cell, '1F497D')
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(14)
    add_run(p, "الجمهورية العراقية", 20, True, C_WHITE)

    add_para(doc, space_before=6, space_after=0)

    t2 = doc.add_table(rows=1, cols=1)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    c2 = t2.cell(0, 0)
    shade_cell(c2, 'DEF0FA')
    p2 = c2.paragraphs[0]
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(20)
    p2.paragraph_format.space_after  = Pt(10)
    add_run(p2, info["عنوان_التقرير"], 22, True, C_DARK)
    p3 = c2.add_paragraph()
    set_rtl(p3)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(20)
    add_run(p3, f"بموجب إطار الشفافية المعزز - اتفاقية باريس", 14, False, C_MED)

    add_para(doc, space_before=20, space_after=4)
    p_info = add_para(doc, f"الجهة المُعِدة: {info['الجهة_المُعِدة']}", 13, False, C_GRAY)
    add_para(doc, f"تاريخ الإصدار: {info['تاريخ_الإصدار']}", 13, False, C_GRAY)
    add_para(doc, f"الفترة المرجعية: {info['سنوات_التقرير']}", 13, False, C_GRAY)
    add_para(doc, f"الإرشادات المعتمدة: {info['إرشادات_IPCC']} | {info['إمكانية_احترارية']}", 13, False, C_GRAY)

    doc.add_page_break()

    # =========================================
    # القسم الأول: الجرد الوطني
    # =========================================
    add_heading(doc, "أولاً: الجرد الوطني لغازات الاحتباس الحراري", 1)
    add_para(doc,
        f"يعرض هذا القسم نتائج الجرد الوطني للعراق للفترة {info['سنوات_التقرير']}، "
        f"المُعَدّ وفق إرشادات {info['إرشادات_IPCC']} باستخدام قيم إمكانية الاحترار "
        f"العالمي {info['إمكانية_احترارية']}.",
        size=13, space_after=8)

    # بطاقات الملخص
    inv = data["الجرد_الوطني"]
    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    headers = ["إجمالي 2021 (Gg CO2eq)", "إجمالي 2020 (Gg CO2eq)", "التغيير منذ 1990"]
    values  = [
        f"{inv['إجمالي_2021_بدون_LULUCF']:,}",
        f"{inv['إجمالي_2020_بدون_LULUCF']:,}",
        f"+{inv['التغيير_منذ_1990_بالمئة']}%"
    ]
    for i, h in enumerate(headers):
        set_cell_text(tbl.cell(0, i), h, 12, True, C_WHITE, WD_ALIGN_PARAGRAPH.CENTER, '1F497D')
    for i, v in enumerate(values):
        color = C_RED if i == 2 else C_DARK
        set_cell_text(tbl.cell(1, i), v, 16, True, color, WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, space_before=8, space_after=4)
    add_heading(doc, "1.1  الانبعاثات حسب القطاع", 2)

    # جدول القطاعات
    cols = ["القطاع", "انبعاثات 2021\n(Gg CO2eq)", "انبعاثات 2020\n(Gg CO2eq)", "الملاحظة"]
    tbl2 = doc.add_table(rows=1 + len(inv["القطاعات"]), cols=4)
    tbl2.style   = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(4.5), Cm(3.5), Cm(3.5), Cm(6)]
    for i, (h, w) in enumerate(zip(cols, widths)):
        c = tbl2.cell(0, i)
        c.width = w
        set_cell_text(c, h, 12, True, C_WHITE, WD_ALIGN_PARAGRAPH.CENTER, '2E75B6')
    for r, sec in enumerate(inv["القطاعات"], 1):
        bg = 'F2F7FD' if r % 2 == 0 else 'FFFFFF'
        set_cell_text(tbl2.cell(r, 0), f"{sec['الرمز']} - {sec['الاسم']}", 12, False, C_TEXT, WD_ALIGN_PARAGRAPH.RIGHT, bg)
        set_cell_text(tbl2.cell(r, 1), f"{sec['2021']:,}", 12, True, C_DARK, WD_ALIGN_PARAGRAPH.CENTER, bg)
        set_cell_text(tbl2.cell(r, 2), f"{sec['2020']:,}", 12, False, C_TEXT, WD_ALIGN_PARAGRAPH.CENTER, bg)
        set_cell_text(tbl2.cell(r, 3), sec["الملاحظة"], 11, False, C_GRAY, WD_ALIGN_PARAGRAPH.RIGHT, bg)

    add_separator(doc)

    # =========================================
    # القسم الثاني: التقدم نحو NDC
    # =========================================
    add_heading(doc, "ثانياً: التقدم في تنفيذ المساهمة الوطنية المحددة (NDC)", 1)
    ndc = data["التقدم_نحو_NDC"]
    add_para(doc,
        f"تتضمن المساهمة الوطنية المحددة للعراق هدفاً {ndc['هدف_NDC']['النوع']}، "
        f"يرمي إلى {ndc['هدف_NDC']['الهدف_غير_المشروط']} بجهود ذاتية، "
        f"و{ndc['هدف_NDC']['الهدف_المشروط']}.",
        size=13, space_after=8)

    add_heading(doc, "2.1  مؤشرات التقدم", 2)

    # جدول المؤشرات
    tbl3 = doc.add_table(rows=1 + len(ndc["مؤشرات_التقدم"]), cols=5)
    tbl3.style = 'Table Grid'
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdrs3 = ["المؤشر", "الأساس 2021", "الوضع الحالي 2023", "الهدف 2030", "الاتجاه"]
    for i, h in enumerate(hdrs3):
        set_cell_text(tbl3.cell(0, i), h, 12, True, C_WHITE, WD_ALIGN_PARAGRAPH.CENTER, '2E75B6')
    for r, ind in enumerate(ndc["مؤشرات_التقدم"], 1):
        bg = 'F2F7FD' if r % 2 == 0 else 'FFFFFF'
        trend_color = C_GREEN if ind["الاتجاه"] == "إيجابي" else C_RED
        set_cell_text(tbl3.cell(r, 0), f"{ind['المؤشر']}\n({ind['الوحدة']})", 11, False, C_TEXT, WD_ALIGN_PARAGRAPH.RIGHT, bg)
        set_cell_text(tbl3.cell(r, 1), str(ind["الأساس_2021"]), 12, False, C_TEXT, WD_ALIGN_PARAGRAPH.CENTER, bg)
        set_cell_text(tbl3.cell(r, 2), str(ind["الحالي_2023"]), 12, True, C_DARK, WD_ALIGN_PARAGRAPH.CENTER, bg)
        set_cell_text(tbl3.cell(r, 3), str(ind["الهدف_2030"]), 12, False, C_TEXT, WD_ALIGN_PARAGRAPH.CENTER, bg)
        set_cell_text(tbl3.cell(r, 4), f"↑ {ind['الاتجاه']}", 12, True, trend_color, WD_ALIGN_PARAGRAPH.CENTER, bg)

    add_separator(doc)

    # =========================================
    # القسم الثالث: السياسات والإجراءات
    # =========================================
    add_heading(doc, "ثالثاً: السياسات والإجراءات المناخية", 1)
    add_para(doc,
        "يستعرض هذا القسم أبرز السياسات والإجراءات الوطنية المنفّذة للحد من انبعاثات "
        "غازات الاحتباس الحراري في القطاعات ذات الأولوية.",
        size=13, space_after=8)

    for pol in data["السياسات_والإجراءات"]:
        add_heading(doc, f"{pol['الرقم']}. {pol['الاسم']}", 2)
        tbl_p = doc.add_table(rows=4, cols=2)
        tbl_p.style = 'Table Grid'
        rows_data = [
            ("القطاع", pol["القطاع"]),
            ("النوع والحالة", f"{pol['النوع']} | {pol['الحالة']}"),
            ("التخفيض المتوقع 2030", pol["التخفيض_2030"]),
            ("الوصف", pol["الوصف"]),
        ]
        for ri, (label, value) in enumerate(rows_data):
            bg = 'EEF4FB' if ri % 2 == 0 else 'FFFFFF'
            set_cell_text(tbl_p.cell(ri, 0), label, 12, True, C_DARK, WD_ALIGN_PARAGRAPH.RIGHT, bg)
            set_cell_text(tbl_p.cell(ri, 1), value, 12, False, C_TEXT, WD_ALIGN_PARAGRAPH.RIGHT, bg)
        add_para(doc, space_before=4, space_after=4)

    add_separator(doc)

    # =========================================
    # القسم الرابع: احتياجات الدعم
    # =========================================
    add_heading(doc, "رابعاً: احتياجات الدعم المالي والتكنولوجي وبناء القدرات", 1)
    sup  = data["احتياجات_الدعم"]
    fin  = sup["التمويل"]

    # بطاقات التمويل
    tbl4 = doc.add_table(rows=2, cols=3)
    tbl4.style = 'Table Grid'
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    fin_headers = ["الاحتياج الإجمالي", "التمويل المُستلَم", "الفجوة التمويلية"]
    fin_values  = [
        f"{fin['الإجمالي_المطلوب_مليار_دولار']} مليار $",
        f"{fin['المُستلَم_مليون_دولار']} مليون $",
        f"{fin['الفجوة_مليار_دولار']} مليار $"
    ]
    fin_colors  = ['1F497D', '37864B', 'C00000']
    for i, (h, v, bg) in enumerate(zip(fin_headers, fin_values, fin_colors)):
        set_cell_text(tbl4.cell(0, i), h, 12, True, C_WHITE, WD_ALIGN_PARAGRAPH.CENTER, bg)
        set_cell_text(tbl4.cell(1, i), v, 15, True, RGBColor(*bytes.fromhex(bg)),
                      WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, space_before=10, space_after=4)
    add_heading(doc, "4.1  الاحتياجات التكنولوجية", 2)
    for item in sup["الاحتياجات_التكنولوجية"]:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, "◀  " + item, 13, False, C_TEXT)

    add_para(doc, space_before=6, space_after=4)
    add_heading(doc, "4.2  بناء القدرات", 2)
    for item in sup["بناء_القدرات"]:
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, "◀  " + item, 13, False, C_TEXT)

    # =========================================
    # التذييل
    # =========================================
    section = doc.sections[-1]
    footer  = section.footer
    ft_p    = footer.paragraphs[0]
    set_rtl(ft_p)
    ft_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(ft_p, f"{info['عنوان_التقرير']}  |  {info['الجهة_المُعِدة']}  |  {info['تاريخ_الإصدار']}",
            10, False, C_GRAY)

    return doc

# ========== تشغيل المولّد ==========
if __name__ == "__main__":
    print("جارٍ توليد تقرير BTR...")
    doc = build_btr(BTR_DATA)

    output = "tools/btr_generator/output_BTR1_العراق.docx"
    doc.save(output)
    print(f"تم إنشاء التقرير: {output}")
    print(f"يمكنك فتحه مباشرة من VS Code أو من مجلد المشروع")
