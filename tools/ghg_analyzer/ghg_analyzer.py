"""
أداة تحليل بيانات الجرد الوطني مع رسوم بيانية عربية
التشغيل: python tools/ghg_analyzer/ghg_analyzer.py
المخرجات: تقرير Word + رسوم PNG
"""

import sys, os, io
sys.path.insert(0, os.path.dirname(__file__))
from ghg_data import (YEARS, SECTORS_DATA, GAS_DATA_2021,
                       ENERGY_SUBSECTORS_2021, REGIONAL_COMPARISON, NDC_TARGETS)

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import arabic_reshaper
from bidi.algorithm import get_display
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches

import datetime, json

# ========== إعداد الخط العربي للرسوم ==========
plt.rcParams['font.family']     = 'DejaVu Sans'
plt.rcParams['axes.unicode_minus'] = False

def ar(text):
    """تحويل النص العربي للعرض الصحيح في matplotlib"""
    return get_display(arabic_reshaper.reshape(str(text)))

# ========== ألوان ==========
PALETTE = ['#1F497D','#2E75B6','#4BACC6','#70AD47','#FFC000','#FF7043','#9E9E9E','#5C4033']
C_DARK  = RGBColor(0x1F, 0x49, 0x7D)
C_MED   = RGBColor(0x2E, 0x75, 0xB6)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY  = RGBColor(0x70, 0x70, 0x70)
C_TEXT  = RGBColor(0x26, 0x26, 0x26)
C_GREEN = RGBColor(0x37, 0x86, 0x44)
C_RED   = RGBColor(0xC0, 0x00, 0x00)
FONT    = "Simplified Arabic"

# ========== دوال الرسوم البيانية ==========
def fig_to_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def chart_trend():
    """رسم الاتجاه الزمني للانبعاثات الكلية"""
    totals = [sum(SECTORS_DATA[s][i] for s in SECTORS_DATA)
              for i in range(len(YEARS))]

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.fill_between(YEARS, totals, alpha=0.15, color='#1F497D')
    ax.plot(YEARS, totals, color='#1F497D', linewidth=2.5, marker='o',
            markersize=5, markevery=2)

    # خط هدف NDC
    ax.axhline(NDC_TARGETS["غير مشروط 2030"], color='#C00000',
               linewidth=1.8, linestyle='--', alpha=0.8)
    ax.axhline(NDC_TARGETS["مشروط 2030"], color='#37864B',
               linewidth=1.8, linestyle='--', alpha=0.8)

    ax.set_xlabel(ar("السنة"), fontsize=12)
    ax.set_ylabel(ar("Gg CO2eq"), fontsize=12)
    ax.set_title(ar("الاتجاه الزمني للانبعاثات الوطنية 2000-2021"), fontsize=14, fontweight='bold')
    ax.set_xticks(YEARS[::2])
    ax.set_xticklabels([str(y) for y in YEARS[::2]], rotation=45)
    ax.yaxis.set_major_formatter(
        matplotlib.ticker.FuncFormatter(lambda x, _: f'{x/1000:.0f}k'))
    ax.legend([
        ar("الإجمالي"),
        ar("هدف غير مشروط 2030"),
        ar("هدف مشروط 2030"),
    ], fontsize=10, loc='upper left')
    ax.grid(axis='y', alpha=0.3)
    ax.spines[['top','right']].set_visible(False)
    fig.tight_layout()
    return fig_to_bytes(fig)

def chart_sectors_stacked():
    """رسم الانبعاثات المكدّسة حسب القطاع"""
    fig, ax = plt.subplots(figsize=(11, 5))
    bottom = np.zeros(len(YEARS))
    for i, (sector, values) in enumerate(SECTORS_DATA.items()):
        ax.bar(YEARS, values, bottom=bottom,
               color=PALETTE[i], label=ar(sector), alpha=0.9, width=0.8)
        bottom += np.array(values)

    ax.set_xlabel(ar("السنة"), fontsize=12)
    ax.set_ylabel(ar("Gg CO2eq"), fontsize=12)
    ax.set_title(ar("الانبعاثات حسب القطاع 2000-2021"), fontsize=14, fontweight='bold')
    ax.set_xticks(YEARS[::3])
    ax.set_xticklabels([str(y) for y in YEARS[::3]], rotation=45)
    ax.yaxis.set_major_formatter(
        matplotlib.ticker.FuncFormatter(lambda x, _: f'{x/1000:.0f}k'))
    ax.legend(fontsize=10, loc='upper left',
              handles=[mpatches.Patch(color=PALETTE[i], label=ar(s))
                       for i, s in enumerate(SECTORS_DATA)])
    ax.grid(axis='y', alpha=0.3)
    ax.spines[['top','right']].set_visible(False)
    fig.tight_layout()
    return fig_to_bytes(fig)

def chart_gases_pie():
    """رسم دائري للانبعاثات حسب الغاز"""
    labels = [ar(k) for k in GAS_DATA_2021]
    values = list(GAS_DATA_2021.values())
    explode = [0.05] * len(values)

    fig, ax = plt.subplots(figsize=(7, 6))
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, autopct='%1.1f%%',
        colors=PALETTE[:len(values)], explode=explode,
        startangle=140, pctdistance=0.82,
        textprops={'fontsize': 11})
    for at in autotexts:
        at.set_fontsize(10)
        at.set_fontweight('bold')
    ax.set_title(ar("توزيع الانبعاثات حسب نوع الغاز - 2021"), fontsize=13, fontweight='bold')
    fig.tight_layout()
    return fig_to_bytes(fig)

def chart_energy_subsectors():
    """رسم أفقي للقطاعات الفرعية للطاقة"""
    labels = [ar(k) for k in ENERGY_SUBSECTORS_2021]
    values = list(ENERGY_SUBSECTORS_2021.values())
    sorted_pairs = sorted(zip(values, labels), reverse=False)
    values, labels = zip(*sorted_pairs)

    fig, ax = plt.subplots(figsize=(9, 5))
    bars = ax.barh(labels, values, color=PALETTE[1], alpha=0.85, height=0.6)
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 500, bar.get_y() + bar.get_height()/2,
                f'{val:,}', va='center', fontsize=10)
    ax.set_xlabel(ar("Gg CO2eq"), fontsize=12)
    ax.set_title(ar("انبعاثات قطاع الطاقة حسب الفئة الفرعية - 2021"),
                 fontsize=13, fontweight='bold')
    ax.xaxis.set_major_formatter(
        matplotlib.ticker.FuncFormatter(lambda x, _: f'{x/1000:.0f}k'))
    ax.grid(axis='x', alpha=0.3)
    ax.spines[['top','right']].set_visible(False)
    fig.tight_layout()
    return fig_to_bytes(fig)

def chart_regional():
    """مقارنة إقليمية"""
    countries = [ar(k) for k in REGIONAL_COMPARISON]
    values    = list(REGIONAL_COMPARISON.values())
    colors    = ['#C00000' if 'العراق' in k else '#2E75B6'
                 for k in REGIONAL_COMPARISON]

    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(countries, values, color=colors, alpha=0.85, width=0.6)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 5000,
                f'{val/1000:.0f}k', ha='center', fontsize=9)
    ax.set_ylabel(ar("Gg CO2eq"), fontsize=12)
    ax.set_title(ar("مقارنة إقليمية للانبعاثات - 2020"), fontsize=13, fontweight='bold')
    ax.yaxis.set_major_formatter(
        matplotlib.ticker.FuncFormatter(lambda x, _: f'{x/1000:.0f}k'))
    ax.grid(axis='y', alpha=0.3)
    ax.spines[['top','right']].set_visible(False)
    fig.tight_layout()
    return fig_to_bytes(fig)

# ========== دوال مساعدة Word ==========
def set_rtl(p):
    p._p.get_or_add_pPr().insert(0, OxmlElement('w:bidi'))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def run(p, text, size=13, bold=False, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color or C_TEXT
    return r

def para(doc, text="", size=13, bold=False, color=None, sb=0, sa=5):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        run(p, text, size, bold, color)
    return p

def heading(doc, text, level=1):
    c = {1: C_DARK, 2: C_MED}
    s = {1: 18,     2: 15}
    p = para(doc, text, s[level], True, c[level], sb=10, sa=4)
    if level == 1:
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:color'), '1F497D')
        pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
    return p

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)

def cell_text(cell, text, size=12, bold=False, color=None,
              align=WD_ALIGN_PARAGRAPH.RIGHT, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = align
    run(p, text, size, bold, color or C_TEXT)
    if bg:
        shade(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def insert_chart(doc, img_bytes, caption, width=15):
    doc.add_picture(img_bytes, width=Cm(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = para(doc, caption, 11, False, C_GRAY, sb=2, sa=8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ========== حسابات تحليلية ==========
def compute_stats():
    totals  = [sum(SECTORS_DATA[s][i] for s in SECTORS_DATA)
               for i in range(len(YEARS))]
    latest  = totals[-1]
    prev    = totals[-2]
    base    = totals[0]
    peak    = max(totals)
    yoy     = ((latest - prev) / prev) * 100
    vs_base = ((latest - base) / base) * 100
    cagr    = ((latest / base) ** (1/(len(YEARS)-1)) - 1) * 100
    dominant= max(SECTORS_DATA, key=lambda s: SECTORS_DATA[s][-1])
    dom_pct = SECTORS_DATA[dominant][-1] / latest * 100
    return {
        "الإجمالي_2021": latest, "الإجمالي_2020": prev,
        "التغيير_السنوي_pct": round(yoy, 2),
        "التغيير_منذ_2000_pct": round(vs_base, 2),
        "CAGR_pct": round(cagr, 2),
        "القمة": peak, "القمة_سنة": YEARS[totals.index(peak)],
        "القطاع_المهيمن": dominant,
        "حصة_القطاع_المهيمن_pct": round(dom_pct, 1),
    }

# ========== بناء التقرير ==========
def build_report():
    doc = Document()
    for s in doc.sections:
        s.page_width    = Cm(21)
        s.page_height   = Cm(29.7)
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.right_margin  = Cm(2.0)
        s.left_margin   = Cm(2.0)
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(13)

    today = datetime.date.today().strftime("%d/%m/%Y")
    stats = compute_stats()

    # === الغلاف ===
    para(doc, sb=20, sa=0)
    t = doc.add_table(1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, '1F497D')
    p = c.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run(p, "تقرير تحليل بيانات الجرد الوطني", 22, True, C_WHITE)
    p2 = c.add_paragraph()
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    run(p2, "غازات الاحتباس الحراري - العراق 2000-2021", 15, False, C_WHITE)

    para(doc, f"تاريخ الإصدار: {today}  |  وزارة البيئة - قسم سياسات المناخ",
         12, False, C_GRAY, sb=8, sa=4)
    para(doc, "الوحدة: Gg CO2eq  |  إرشادات IPCC 2006  |  إمكانية الاحترار AR5",
         12, False, C_GRAY, sb=0, sa=4)
    doc.add_page_break()

    # === القسم الأول: المؤشرات الرئيسية ===
    heading(doc, "أولاً: المؤشرات الرئيسية", 1)

    kpis = [
        ("إجمالي 2021",         f"{stats['الإجمالي_2021']:,} Gg CO2eq", "1F497D"),
        ("التغيير السنوي",       f"{stats['التغيير_السنوي_pct']:+.1f}%",
         "C00000" if stats['التغيير_السنوي_pct'] > 0 else "37864B"),
        ("معدل النمو السنوي",    f"{stats['CAGR_pct']:.2f}% / سنة",    "BF8F00"),
        ("التغيير منذ 2000",     f"{stats['التغيير_منذ_2000_pct']:+.1f}%","C00000"),
        ("القطاع المهيمن",       f"{stats['القطاع_المهيمن']} ({stats['حصة_القطاع_المهيمن_pct']:.0f}%)", "2E75B6"),
        ("ذروة الانبعاثات",      f"{stats['القمة']:,} ({stats['القمة_سنة']})", "7030A0"),
    ]
    tbl_k = doc.add_table(2, 3)
    tbl_k.style = "Table Grid"
    tbl_k.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(3):
        cell_text(tbl_k.cell(0, i), kpis[i][0],   12, True,  C_WHITE,
                  WD_ALIGN_PARAGRAPH.CENTER, kpis[i][2])
        cell_text(tbl_k.cell(1, i), kpis[i][1],   15, True,
                  RGBColor(*bytes.fromhex(kpis[i][2])),
                  WD_ALIGN_PARAGRAPH.CENTER)

    para(doc, sb=6, sa=0)
    tbl_k2 = doc.add_table(2, 3)
    tbl_k2.style = "Table Grid"
    tbl_k2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(3):
        cell_text(tbl_k2.cell(0, i), kpis[i+3][0], 12, True, C_WHITE,
                  WD_ALIGN_PARAGRAPH.CENTER, kpis[i+3][2])
        cell_text(tbl_k2.cell(1, i), kpis[i+3][1], 15, True,
                  RGBColor(*bytes.fromhex(kpis[i+3][2])),
                  WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # === القسم الثاني: الاتجاه الزمني ===
    heading(doc, "ثانياً: الاتجاه الزمني للانبعاثات", 1)
    para(doc,
         f"شهدت الانبعاثات الوطنية نمواً بمعدل {stats['CAGR_pct']:.2f}% سنوياً خلال الفترة "
         f"2000-2021، مرتفعةً من {stats['الإجمالي_2020']//1000:.0f} ألف إلى "
         f"{stats['الإجمالي_2021']//1000:.0f} ألف Gg CO2eq. "
         f"ويظهر الرسم البياني مقارنة الاتجاه بأهداف المساهمة الوطنية لعام 2030.",
         13, sa=8)
    insert_chart(doc, chart_trend(),
                 "شكل (1): الاتجاه الزمني للانبعاثات الوطنية مقارنةً بأهداف NDC")

    # === القسم الثالث: تحليل القطاعات ===
    heading(doc, "ثالثاً: تحليل الانبعاثات حسب القطاع", 1)
    para(doc,
         f"يهيمن قطاع {stats['القطاع_المهيمن']} على الانبعاثات الوطنية بحصة "
         f"{stats['حصة_القطاع_المهيمن_pct']:.0f}% من الإجمالي عام 2021، "
         f"مما يجعله القطاع ذا الأولوية القصوى في خطط التخفيف.",
         13, sa=8)
    insert_chart(doc, chart_sectors_stacked(),
                 "شكل (2): الانبعاثات المكدّسة حسب القطاع 2000-2021")

    # جدول القطاعات
    heading(doc, "3.1  مقارنة القطاعات 2020 و 2021", 2)
    hdrs = ["القطاع", "2020 (Gg CO2eq)", "2021 (Gg CO2eq)", "التغيير", "الحصة 2021"]
    widths = [4.5, 3.5, 3.5, 3.0, 3.0]
    tbl_s = doc.add_table(1 + len(SECTORS_DATA), 5)
    tbl_s.style = "Table Grid"
    tbl_s.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_2021 = sum(SECTORS_DATA[s][-1] for s in SECTORS_DATA)
    for i, (h, w) in enumerate(zip(hdrs, widths)):
        c = tbl_s.cell(0, i)
        c.width = Cm(w)
        cell_text(c, h, 12, True, C_WHITE, WD_ALIGN_PARAGRAPH.CENTER, '2E75B6')
    for ri, (sec, vals) in enumerate(SECTORS_DATA.items(), 1):
        bg = 'F2F7FD' if ri % 2 == 0 else 'FFFFFF'
        v20, v21 = vals[-2], vals[-1]
        chg = ((v21 - v20) / v20) * 100
        pct = v21 / total_2021 * 100
        chg_color = 'C00000' if chg > 0 else '37864B'
        cell_text(tbl_s.cell(ri, 0), sec, 12, False, C_TEXT, WD_ALIGN_PARAGRAPH.RIGHT, bg)
        cell_text(tbl_s.cell(ri, 1), f"{v20:,}", 12, False, C_TEXT, WD_ALIGN_PARAGRAPH.CENTER, bg)
        cell_text(tbl_s.cell(ri, 2), f"{v21:,}", 12, True,  C_DARK, WD_ALIGN_PARAGRAPH.CENTER, bg)
        cell_text(tbl_s.cell(ri, 3), f"{chg:+.1f}%", 12, True,
                  RGBColor(*bytes.fromhex(chg_color)), WD_ALIGN_PARAGRAPH.CENTER, bg)
        cell_text(tbl_s.cell(ri, 4), f"{pct:.1f}%", 12, False, C_TEXT, WD_ALIGN_PARAGRAPH.CENTER, bg)

    doc.add_page_break()

    # === القسم الرابع: تحليل الغازات ===
    heading(doc, "رابعاً: توزيع الانبعاثات حسب نوع الغاز", 1)
    insert_chart(doc, chart_gases_pie(),
                 "شكل (3): توزيع الانبعاثات حسب نوع الغاز - 2021", width=12)

    # === القسم الخامس: قطاع الطاقة ===
    heading(doc, "خامساً: تحليل قطاع الطاقة", 1)
    para(doc, "يُعدّ قطاع الطاقة المصدر الرئيسي للانبعاثات، ويوضح الرسم التالي "
         "التوزيع الداخلي للانبعاثات بين فئاته الفرعية المختلفة.", 13, sa=8)
    insert_chart(doc, chart_energy_subsectors(),
                 "شكل (4): انبعاثات قطاع الطاقة حسب الفئة الفرعية - 2021")

    doc.add_page_break()

    # === القسم السادس: المقارنة الإقليمية ===
    heading(doc, "سادساً: المقارنة الإقليمية", 1)
    para(doc, "يُصنَّف العراق في المرتبة الثالثة إقليمياً من حيث إجمالي الانبعاثات. "
         "تُبرز المقارنة الإقليمية الوزن النسبي للعراق وأهمية مشاركته في "
         "مفاوضات المناخ الدولية.", 13, sa=8)
    insert_chart(doc, chart_regional(),
                 "شكل (5): مقارنة إقليمية للانبعاثات - 2020")

    # === التذييل ===
    footer = doc.sections[-1].footer.paragraphs[0]
    set_rtl(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(footer,
        f"تقرير تحليل الجرد الوطني - العراق  |  {today}  |  وزارة البيئة",
        10, False, C_GRAY)

    return doc, stats

# ========== التشغيل ==========
if __name__ == "__main__":
    print("جارٍ تحليل بيانات الجرد وإنشاء الرسوم البيانية...")

    doc, stats = build_report()

    out_docx = "tools/ghg_analyzer/output_GHG_تقرير_التحليل.docx"
    doc.save(out_docx)
    print(f"✓ تم إنشاء التقرير: {out_docx}")

    out_json = "tools/ghg_analyzer/output_GHG_إحصاءات.json"
    with open(out_json, 'w', encoding='utf-8') as f:
        json.dump(stats, f, ensure_ascii=False, indent=2)
    print(f"✓ تم حفظ الإحصاءات: {out_json}")

    print("\n" + "="*55)
    print("الإحصاءات الرئيسية")
    print("="*55)
    for k, v in stats.items():
        print(f"  {k}: {v}")
