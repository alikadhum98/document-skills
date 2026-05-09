"""
أداة تقييم NDC على مسارات Navigator السبعة - العراق
التشغيل: python tools/ndc_navigator/ndc_navigator.py
"""

import sys, os, io, json, datetime
sys.path.insert(0, os.path.dirname(__file__))
from ndc_data import PARTY, ASSESSMENT_DATE, NDC_VERSION, ROUTES, PRIORITY_ACTIONS

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import arabic_reshaper
from bidi.algorithm import get_display

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['axes.unicode_minus'] = False

def ar(t): return get_display(arabic_reshaper.reshape(str(t)))

C_DARK  = RGBColor(0x1F,0x49,0x7D)
C_MED   = RGBColor(0x2E,0x75,0xB6)
C_GREEN = RGBColor(0x37,0x86,0x44)
C_RED   = RGBColor(0xC0,0x00,0x00)
C_AMBER = RGBColor(0xBF,0x8F,0x00)
C_WHITE = RGBColor(0xFF,0xFF,0xFF)
C_GRAY  = RGBColor(0x70,0x70,0x70)
C_TEXT  = RGBColor(0x26,0x26,0x26)
FONT    = "Simplified Arabic"

STATUS_BG = {
    "جزئي":          "2E75B6",
    "يحتاج تطوير":   "BF8F00",
    "غائب":          "C00000",
    "ممتاز":         "37864B",
}

PRIORITY_BG = {"عالية":"C00000","متوسطة":"BF8F00","منخفضة":"37864B"}

def fig_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf,format='png',dpi=150,bbox_inches='tight',facecolor='white')
    buf.seek(0); plt.close(fig); return buf

def chart_radar():
    labels  = [ar(r["الاسم"][:12]) for r in ROUTES]
    scores  = [r["الدرجة"] for r in ROUTES]
    n       = len(labels)
    angles  = [i * 2 * np.pi / n for i in range(n)] + [0]
    scores_ = scores + [scores[0]]

    fig, ax = plt.subplots(figsize=(8,8), subplot_kw=dict(polar=True))
    ax.fill(angles[:-1], scores, alpha=0.25, color='#2E75B6')
    ax.plot(angles, scores_, color='#1F497D', linewidth=2.5, marker='o', markersize=7)
    ax.fill(angles[:-1], [5]*n, alpha=0.05, color='#E8F0F7')

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_yticks([1,2,3,4,5])
    ax.set_yticklabels(['1','2','3','4','5'], fontsize=9)
    ax.set_ylim(0,5)
    ax.set_title(ar("تقييم NDC على مسارات Navigator السبعة"),
                 fontsize=13, fontweight='bold', pad=20)
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    return fig_bytes(fig)

def chart_bars():
    labels = [ar(r["الرمز"]+": "+r["الاسم"][:15]) for r in ROUTES]
    scores = [r["الدرجة"] for r in ROUTES]
    colors = ['#37864B' if s>=4 else '#BF8F00' if s==3 else '#C00000' for s in scores]

    fig, ax = plt.subplots(figsize=(11,5))
    bars = ax.bar(labels, scores, color=colors, alpha=0.85, width=0.6)
    ax.axhline(3, color='#BF8F00', linewidth=1.5, linestyle='--', alpha=0.7,
               label=ar("الحد المقبول"))
    ax.axhline(4, color='#37864B', linewidth=1.5, linestyle='--', alpha=0.7,
               label=ar("الهدف المطلوب"))
    for bar, score in zip(bars, scores):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.05,
                f'{score}/5', ha='center', fontsize=11, fontweight='bold')
    ax.set_ylim(0, 6)
    ax.set_ylabel(ar("الدرجة من 5"), fontsize=11)
    ax.set_title(ar("درجات التقييم على كل مسار"), fontsize=13, fontweight='bold')
    ax.legend(fontsize=10)
    ax.grid(axis='y', alpha=0.3)
    ax.spines[['top','right']].set_visible(False)
    fig.tight_layout()
    return fig_bytes(fig)

# ===== دوال Word =====
def set_rtl(p):
    p._p.get_or_add_pPr().insert(0, OxmlElement('w:bidi'))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def run(p, t, size=13, bold=False, color=None):
    r = p.add_run(t); r.font.name=FONT; r.font.size=Pt(size)
    r.font.bold=bold; r.font.color.rgb=color or C_TEXT; return r

def para(doc,t="",size=13,bold=False,color=None,sb=0,sa=5):
    p=doc.add_paragraph(); set_rtl(p)
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if t: run(p,t,size,bold,color)
    return p

def heading(doc,t,level=1):
    c={1:C_DARK,2:C_MED}; s={1:18,2:15}
    p=para(doc,t,s[level],True,c[level],sb=10,sa=4)
    if level==1:
        pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:color'),'1F497D'); pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
    return p

def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:fill'),hx); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def ctxt(cell,t,size=12,bold=False,color=None,
         align=WD_ALIGN_PARAGRAPH.RIGHT,bg=None):
    cell.text=""
    p=cell.paragraphs[0]; set_rtl(p); p.alignment=align
    run(p,t,size,bold,color or C_TEXT)
    if bg: shade(cell,bg)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

def insert_chart(doc,buf,caption,width=15):
    doc.add_picture(buf,width=Cm(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=para(doc,caption,11,False,C_GRAY,sb=2,sa=8)
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER

def build_report():
    doc=Document(); today=datetime.date.today().strftime("%d/%m/%Y")
    avg=sum(r["الدرجة"] for r in ROUTES)/len(ROUTES)

    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.right_margin=Cm(2.0); sec.left_margin=Cm(2.0)
    doc.styles['Normal'].font.name=FONT
    doc.styles['Normal'].font.size=Pt(13)

    # الغلاف
    para(doc,sb=20,sa=0)
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,'1F497D')
    p=c.paragraphs[0]; set_rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(6)
    run(p,"تقرير تقييم المساهمة الوطنية المحددة",22,True,C_WHITE)
    p2=c.add_paragraph(); set_rtl(p2); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after=Pt(16)
    run(p2,f"NDC Navigator - {PARTY} | {NDC_VERSION}",15,False,C_WHITE)
    para(doc,f"تاريخ التقييم: {today}  |  المرجع: NDC 3.0 Navigator",12,False,C_GRAY,sb=8,sa=4)
    doc.add_page_break()

    # ملخص الدرجات
    heading(doc,"أولاً: ملخص نتائج التقييم",1)
    avg_color = "37864B" if avg>=4 else "BF8F00" if avg>=3 else "C00000"
    kpis=[
        ("المتوسط الكلي",f"{avg:.1f}/5.0",avg_color),
        ("مسارات عالية الأولوية",
         str(sum(1 for r in ROUTES if r["الأولوية"]=="عالية")),"C00000"),
        ("مسارات غائبة",
         str(sum(1 for r in ROUTES if r["الحالة"]=="غائب")),"7030A0"),
    ]
    tbl_k=doc.add_table(2,3); tbl_k.style="Table Grid"
    tbl_k.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(label,val,bg) in enumerate(kpis):
        ctxt(tbl_k.cell(0,i),label,12,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl_k.cell(1,i),val,18,True,
             RGBColor(*bytes.fromhex(bg)),WD_ALIGN_PARAGRAPH.CENTER)

    para(doc,sb=8,sa=4)
    insert_chart(doc,chart_radar(),"شكل (1): مخطط الرادار - مستوى الأداء على المسارات السبعة",width=13)
    insert_chart(doc,chart_bars(),"شكل (2): الدرجات التفصيلية على كل مسار",width=15)
    doc.add_page_break()

    # تفاصيل كل مسار
    heading(doc,"ثانياً: التقييم التفصيلي لكل مسار",1)
    for r in ROUTES:
        sbg=STATUS_BG.get(r["الحالة"],"888888")
        pbg=PRIORITY_BG.get(r["الأولوية"],"888888")
        heading(doc,f"{r['الرمز']}: {r['الاسم']}",2)

        # بطاقة المسار
        tbl=doc.add_table(1,4); tbl.style="Table Grid"
        tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        ctxt(tbl.cell(0,0),"الدرجة",11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,avg_color)
        ctxt(tbl.cell(0,1),f"{r['الدرجة']}/5",16,True,
             RGBColor(*bytes.fromhex(avg_color)),WD_ALIGN_PARAGRAPH.CENTER)
        ctxt(tbl.cell(0,2),"الحالة",11,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,sbg)
        ctxt(tbl.cell(0,3),r["الحالة"],13,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,sbg)

        para(doc,sb=4,sa=2)
        para(doc,r["الوصف"],12,False,C_GRAY,sa=6)

        # الفجوات
        para(doc,"الفجوات الرئيسية:",13,True,C_RED,sb=4,sa=2)
        for gap in r["الفجوات"]:
            p=doc.add_paragraph(); set_rtl(p)
            p.paragraph_format.space_after=Pt(3)
            run(p,"✗  "+gap,12,False,C_TEXT)

        # التوصيات
        para(doc,"التوصيات:",13,True,C_GREEN,sb=4,sa=2)
        for rec in r["التوصيات"]:
            p=doc.add_paragraph(); set_rtl(p)
            p.paragraph_format.space_after=Pt(3)
            run(p,"✓  "+rec,12,False,C_TEXT)

        # الموعد والأولوية
        p=doc.add_paragraph(); set_rtl(p)
        p.paragraph_format.space_after=Pt(8)
        run(p,f"الأولوية: {r['الأولوية']}  |  الموعد المقترح: {r['الموعد']}",
            12,True,RGBColor(*bytes.fromhex(pbg)))

    doc.add_page_break()

    # الأولويات الفورية
    heading(doc,"ثالثاً: الإجراءات ذات الأولوية الفورية",1)
    para(doc,"بناءً على نتائج التقييم، تُوصى الجهات المعنية باتخاذ الإجراءات "
         "التالية قبل تقديم NDC 3.0:",13,sa=8)
    for i,act in enumerate(PRIORITY_ACTIONS,1):
        p=doc.add_paragraph(); set_rtl(p)
        p.paragraph_format.space_after=Pt(5)
        run(p,f"{i}.  ",13,True,C_MED)
        run(p,act,13,False,C_TEXT)

    # التذييل
    footer=doc.sections[-1].footer.paragraphs[0]
    set_rtl(footer); footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run(footer,f"تقرير NDC Navigator - {PARTY}  |  {today}",10,False,C_GRAY)

    return doc, avg

if __name__=="__main__":
    print("جارٍ إنشاء تقرير تقييم NDC...")
    doc, avg = build_report()
    out="tools/ndc_navigator/output_NDC_تقرير_التقييم.docx"
    doc.save(out)
    print(f"✓ تم إنشاء التقرير: {out}")
    print(f"  المتوسط الكلي: {avg:.1f}/5.0")
    high=sum(1 for r in ROUTES if r["الأولوية"]=="عالية")
    print(f"  مسارات عالية الأولوية: {high}")
