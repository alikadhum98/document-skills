"""
أداة تتبع مفاوضات المادة السادسة وحساب ITMOs - العراق
التشغيل: python tools/article6_tracker/a6_tracker.py
المخرجات: تقرير Word + ملف JSON
"""

import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from a6_data import IRAQ_POSITION, DEALS, NEGOTIATION_SESSIONS, ACCOUNTING_PARAMS

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, datetime

# ========== ألوان ==========
C_DARK  = RGBColor(0x1F, 0x49, 0x7D)
C_MED   = RGBColor(0x2E, 0x75, 0xB6)
C_GREEN = RGBColor(0x37, 0x86, 0x44)
C_RED   = RGBColor(0xC0, 0x00, 0x00)
C_AMBER = RGBColor(0xBF, 0x8F, 0x00)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY  = RGBColor(0x70, 0x70, 0x70)
C_TEXT  = RGBColor(0x26, 0x26, 0x26)
FONT    = "Simplified Arabic"

STATUS_COLORS = {
    "نشط":           "37864B",
    "قيد التفاوض":   "2E75B6",
    "مُعلَّق":        "BF8F00",
    "مكتمل":         "1F497D",
    "ملغى":          "C00000",
}

PRIORITY_COLORS = {
    "عالية":   "C00000",
    "متوسطة":  "BF8F00",
    "منخفضة":  "37864B",
}

# ========== دوال مساعدة ==========
def set_rtl(p):
    p._p.get_or_add_pPr().insert(0, OxmlElement('w:bidi'))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def run(p, text, size=13, bold=False, color=None, italic=False):
    r = p.add_run(text)
    r.font.name    = FONT
    r.font.size    = Pt(size)
    r.font.bold    = bold
    r.font.italic  = italic
    r.font.color.rgb = color or C_TEXT
    return r

def para(doc, text="", size=13, bold=False, color=None,
         sb=0, sa=5, italic=False):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        run(p, text, size, bold, color, italic)
    return p

def heading(doc, text, level=1):
    c = {1: C_DARK, 2: C_MED, 3: C_TEXT}
    s = {1: 18,     2: 15,    3: 13}
    p = para(doc, text, s[level], True, c[level], sb=10, sa=4)
    if level == 1:
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:color'), '1F497D')
        pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
    return p

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)

def cell_text(cell, text, size=12, bold=False, color=None,
              align=WD_ALIGN_PARAGRAPH.RIGHT, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = align
    run(p, text, size, bold, color or C_TEXT)
    if bg:
        shade(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def make_table(doc, headers, rows_data, col_widths,
               header_bg="2E75B6", alt_bg="F2F7FD"):
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        c = tbl.cell(0, i)
        c.width = Cm(w)
        cell_text(c, h, 12, True, C_WHITE,
                  WD_ALIGN_PARAGRAPH.CENTER, header_bg)
    for ri, row in enumerate(rows_data):
        bg = alt_bg if ri % 2 == 0 else "FFFFFF"
        for ci, (val, align, bold, color) in enumerate(row):
            cell_text(tbl.cell(ri+1, ci), val, 12, bold,
                      color, align, bg)
    return tbl

# ========== حسابات ITMOs ==========
def calc_deal(deal, params):
    """حساب الإيرادات والتعديلات المقابلة لصفقة واحدة"""
    itmos_total  = deal["ITMOs_المقترحة"] * deal["المدة_سنوات"]
    revenue_total= deal["العائد_السنوي"]  * deal["المدة_سنوات"]
    mech = deal["نوع_الآلية"]
    share_pct = (params["حصة_صندوق_التكيف_6.4_pct"]
                 if "6.4" in mech
                 else params["حصة_صندوق_التكيف_6.2_pct"])
    adaptation_share = revenue_total * share_pct / 100
    net_revenue      = revenue_total - adaptation_share

    # التعديل المقابل: إضافة الـ ITMOs على انبعاثات العراق
    adj_emissions = IRAQ_POSITION["انبعاثات_2021"] + itmos_total
    compliant     = adj_emissions <= IRAQ_POSITION["هدف_NDC_2030"]

    return {
        "ITMOs_الإجمالية_tCO2eq":    itmos_total,
        "الإيرادات_الإجمالية_USD":    revenue_total,
        "حصة_صندوق_التكيف_USD":      adaptation_share,
        "الإيرادات_الصافية_USD":      net_revenue,
        "الانبعاثات_بعد_التعديل":     adj_emissions,
        "امتثال_NDC":                 compliant,
        "تحذير_NDC":                  not compliant,
    }

def calc_portfolio(deals):
    """حساب المحفظة الكاملة"""
    total_itmos   = 0
    total_revenue = 0
    total_adapt   = 0
    warnings      = []
    for d in deals:
        if d["الحالة"] in ["نشط", "قيد التفاوض"]:
            c = calc_deal(d, ACCOUNTING_PARAMS)
            total_itmos   += c["ITMOs_الإجمالية_tCO2eq"]
            total_revenue += c["الإيرادات_الإجمالية_USD"]
            total_adapt   += c["حصة_صندوق_التكيف_USD"]
            if c["تحذير_NDC"]:
                warnings.append(d["معرّف"])
    return {
        "إجمالي_ITMOs": total_itmos,
        "إجمالي_الإيرادات": total_revenue,
        "إجمالي_التكيف": total_adapt,
        "تحذيرات_NDC": warnings,
        "عدد_الصفقات_النشطة": sum(
            1 for d in deals if d["الحالة"] in ["نشط", "قيد التفاوض"]),
    }

# ========== بناء تقرير Word ==========
def build_report(deals, sessions, position, params):
    doc = Document()
    for s in doc.sections:
        s.page_width    = Cm(21)
        s.page_height   = Cm(29.7)
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.right_margin  = Cm(2.5)
        s.left_margin   = Cm(2.5)
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(13)

    today = datetime.date.today().strftime("%d/%m/%Y")

    # === الغلاف ===
    para(doc, sb=20, sa=0)
    t = doc.add_table(1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, '1F497D')
    p = c.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(16)
    run(p, "تقرير متابعة مفاوضات المادة السادسة", 20, True, C_WHITE)

    para(doc, sb=6, sa=0)
    t2 = doc.add_table(1, 1)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    c2 = t2.cell(0, 0)
    shade(c2, 'DEF0FA')
    p2 = c2.paragraphs[0]
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(12)
    p2.paragraph_format.space_after  = Pt(12)
    run(p2, "اتفاقية باريس - العراق", 16, True, C_DARK)
    p3 = c2.add_paragraph()
    set_rtl(p3)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(12)
    run(p3, f"تاريخ التقرير: {today}  |  مجموعة التفاوض: {position['مجموعة_التفاوض']}", 12, False, C_GRAY)

    doc.add_page_break()

    # === القسم الأول: ملخص المحفظة ===
    portfolio = calc_portfolio(deals)
    heading(doc, "أولاً: ملخص محفظة ITMOs", 1)

    # بطاقات الملخص
    cards = [
        ("عدد الصفقات النشطة",   str(portfolio["عدد_الصفقات_النشطة"]),      "2E75B6"),
        ("إجمالي ITMOs (tCO2eq)", f"{portfolio['إجمالي_ITMOs']:,.0f}",       "1F497D"),
        ("إجمالي الإيرادات (USD)",f"${portfolio['إجمالي_الإيرادات']:,.0f}", "37864B"),
        ("حصة صندوق التكيف",     f"${portfolio['إجمالي_التكيف']:,.0f}",     "BF8F00"),
    ]
    t3 = doc.add_table(2, 4)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3.style = "Table Grid"
    for i, (label, value, bg) in enumerate(cards):
        cell_text(t3.cell(0, i), label, 11, True, C_WHITE,
                  WD_ALIGN_PARAGRAPH.CENTER, bg)
        cell_text(t3.cell(1, i), value, 15, True,
                  RGBColor(*bytes.fromhex(bg)),
                  WD_ALIGN_PARAGRAPH.CENTER)

    if portfolio["تحذيرات_NDC"]:
        para(doc, sb=8, sa=0)
        p_warn = para(doc,
            f"⚠ تحذير: الصفقات التالية قد تتجاوز هدف NDC عند الجمع: "
            f"{', '.join(portfolio['تحذيرات_NDC'])}",
            12, True, C_RED, sb=4, sa=4)
    else:
        para(doc, "✓ جميع الصفقات ضمن حدود هدف NDC", 12, True, C_GREEN, sb=6, sa=4)

    # === القسم الثاني: الصفقات التفصيلية ===
    heading(doc, "ثانياً: الصفقات والمفاوضات الجارية", 1)

    for deal in deals:
        heading(doc, f"{deal['معرّف']} - {deal['الدولة_الشريكة']}", 2)
        calc = calc_deal(deal, params)
        status_bg = STATUS_COLORS.get(deal["الحالة"], "888888")

        info_rows = [
            ("الحالة",         deal["الحالة"],         status_bg),
            ("نوع الآلية",     deal["نوع_الآلية"],     None),
            ("القطاع",         deal["القطاع"],          None),
            ("النشاط",         deal["النشاط"],          None),
            ("مدة الصفقة",     f"{deal['المدة_سنوات']} سنة",   None),
            ("سعر الطن",       f"${deal['سعر_الطن_دولار']} / tCO2eq", None),
            ("ITMOs سنوية",    f"{deal['ITMOs_المقترحة']:,} tCO2eq", None),
            ("إجمالي ITMOs",   f"{calc['ITMOs_الإجمالية_tCO2eq']:,} tCO2eq", None),
            ("الإيرادات السنوية", f"${deal['العائد_السنوي']:,}", None),
            ("إجمالي الإيرادات", f"${calc['الإيرادات_الإجمالية_USD']:,}", None),
            ("حصة صندوق التكيف", f"${calc['حصة_صندوق_التكيف_USD']:,}", None),
            ("الإيرادات الصافية", f"${calc['الإيرادات_الصافية_USD']:,}", None),
            ("التحقق",         deal["التحقق"],          None),
            ("امتثال NDC",     "✓ ممتثل" if calc["امتثال_NDC"] else "✗ يتجاوز الهدف",
             "37864B" if calc["امتثال_NDC"] else "C00000"),
            ("ملاحظات",        deal["ملاحظات"],         None),
        ]

        tbl_d = doc.add_table(len(info_rows), 2)
        tbl_d.style = "Table Grid"
        tbl_d.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, (label, value, bg) in enumerate(info_rows):
            row_bg = "EEF4FB" if ri % 2 == 0 else "FFFFFF"
            cell_text(tbl_d.cell(ri, 0), label, 12, True, C_DARK,
                      WD_ALIGN_PARAGRAPH.RIGHT, row_bg)
            val_bg = bg if bg else row_bg
            val_color = C_WHITE if bg else C_TEXT
            val_bold  = bool(bg)
            cell_text(tbl_d.cell(ri, 1), value, 12, val_bold,
                      val_color, WD_ALIGN_PARAGRAPH.RIGHT, val_bg)

        para(doc, sb=6, sa=2)

    # === القسم الثالث: جلسات المفاوضات ===
    heading(doc, "ثالثاً: سجل جلسات المفاوضات", 1)
    hdrs = ["الجلسة", "التاريخ", "الموضوع", "موقف العراق", "النتيجة", "الأولوية"]
    widths = [3.2, 2.0, 3.5, 4.0, 3.5, 2.0]
    rows = []
    for s in sessions:
        pr_color = RGBColor(*bytes.fromhex(
            PRIORITY_COLORS.get(s["الأولوية"], "888888")))
        rows.append([
            (s["الجلسة"],       WD_ALIGN_PARAGRAPH.RIGHT, True,  C_DARK),
            (s["التاريخ"],      WD_ALIGN_PARAGRAPH.CENTER,False, C_TEXT),
            (s["الموضوع"],      WD_ALIGN_PARAGRAPH.RIGHT, False, C_TEXT),
            (s["موقف_العراق"],  WD_ALIGN_PARAGRAPH.RIGHT, False, C_TEXT),
            (s["النتيجة"],      WD_ALIGN_PARAGRAPH.RIGHT, False, C_TEXT),
            (s["الأولوية"],     WD_ALIGN_PARAGRAPH.CENTER,True,  pr_color),
        ])
    make_table(doc, hdrs, rows, widths)

    # === القسم الرابع: المواقف التفاوضية ===
    heading(doc, "رابعاً: المواقف التفاوضية الرئيسية للعراق", 1)
    for i, pos in enumerate(position["مواقف_رئيسية"], 1):
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_after = Pt(4)
        run(p, f"{i}.  ", 13, True, C_MED)
        run(p, pos, 13, False, C_TEXT)

    # === التذييل ===
    footer = doc.sections[-1].footer.paragraphs[0]
    set_rtl(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(footer,
        f"تقرير متابعة المادة السادسة - العراق  |  {today}  |  سري للاستخدام الرسمي",
        10, False, C_GRAY)

    return doc, portfolio

# ========== تشغيل الأداة ==========
if __name__ == "__main__":
    print("جارٍ إنشاء تقرير المادة السادسة...")

    doc, portfolio = build_report(
        DEALS, NEGOTIATION_SESSIONS, IRAQ_POSITION, ACCOUNTING_PARAMS)

    out_docx = "tools/article6_tracker/output_A6_تقرير_المفاوضات.docx"
    doc.save(out_docx)
    print(f"✓ تم إنشاء التقرير: {out_docx}")

    # حفظ البيانات JSON
    out_json = "tools/article6_tracker/output_A6_محفظة_ITMOs.json"
    result = {
        "تاريخ_التقرير": str(datetime.date.today()),
        "ملخص_المحفظة":  portfolio,
        "تفاصيل_الصفقات": [
            {**d, "حسابات": calc_deal(d, ACCOUNTING_PARAMS)}
            for d in DEALS
        ],
    }
    with open(out_json, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f"✓ تم حفظ البيانات: {out_json}")

    # طباعة ملخص في الـ Terminal
    print("\n" + "="*55)
    print("ملخص المحفظة")
    print("="*55)
    print(f"  عدد الصفقات النشطة:   {portfolio['عدد_الصفقات_النشطة']}")
    print(f"  إجمالي ITMOs:         {portfolio['إجمالي_ITMOs']:,} tCO2eq")
    print(f"  إجمالي الإيرادات:     ${portfolio['إجمالي_الإيرادات']:,}")
    print(f"  حصة صندوق التكيف:    ${portfolio['إجمالي_التكيف']:,}")
    if portfolio["تحذيرات_NDC"]:
        print(f"  ⚠ تحذيرات NDC:       {portfolio['تحذيرات_NDC']}")
    else:
        print("  ✓ جميع الصفقات ضمن هدف NDC")
