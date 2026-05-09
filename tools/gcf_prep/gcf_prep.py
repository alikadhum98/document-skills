"""
مولّد وثيقة مفهوم مشروع GCF (Concept Note)
التشغيل: python tools/gcf_prep/gcf_prep.py
"""

import sys, os, json, datetime
sys.path.insert(0, os.path.dirname(__file__))
from gcf_data import PROJECT

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_DARK=RGBColor(0x1F,0x49,0x7D); C_MED=RGBColor(0x2E,0x75,0xB6)
C_GREEN=RGBColor(0x37,0x86,0x44); C_RED=RGBColor(0xC0,0x00,0x00)
C_WHITE=RGBColor(0xFF,0xFF,0xFF); C_GRAY=RGBColor(0x70,0x70,0x70)
C_TEXT=RGBColor(0x26,0x26,0x26); FONT="Simplified Arabic"

def set_rtl(p):
    p._p.get_or_add_pPr().insert(0,OxmlElement('w:bidi'))
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT

def run(p,t,size=13,bold=False,color=None):
    r=p.add_run(t); r.font.name=FONT; r.font.size=Pt(size)
    r.font.bold=bold; r.font.color.rgb=color or C_TEXT; return r

def para(doc,t="",size=13,bold=False,color=None,sb=0,sa=5):
    p=doc.add_paragraph(); set_rtl(p)
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if t: run(p,t,size,bold,color)
    return p

def heading(doc,t,level=1):
    c={1:C_DARK,2:C_MED,3:C_TEXT}; s={1:18,2:15,3:13}
    p=para(doc,t,s[level],True,c[level],sb=10,sa=4)
    if level==1:
        pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:color'),'1F497D'); pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
    return p

def shade(cell,hx):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:fill'),hx); shd.set(qn('w:val'),'clear'); tcPr.append(shd)

def ctxt(cell,t,size=12,bold=False,color=None,
         align=WD_ALIGN_PARAGRAPH.RIGHT,bg=None):
    cell.text=""; p=cell.paragraphs[0]; set_rtl(p); p.alignment=align
    run(p,t,size,bold,color or C_TEXT)
    if bg: shade(cell,bg)
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

def info_table(doc, rows):
    tbl=doc.add_table(len(rows),2); tbl.style="Table Grid"
    tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ri,(l,v) in enumerate(rows):
        bg="EEF4FB" if ri%2==0 else "FFFFFF"
        ctxt(tbl.cell(ri,0),l,12,True,C_DARK,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl.cell(ri,1),v,12,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)
    para(doc,sb=4,sa=4)

def build_concept_note():
    doc=Document(); today=datetime.date.today().strftime("%d/%m/%Y")
    total=PROJECT["تمويل_GCF_مليون"]+PROJECT["تمويل_مشترك_مليون"]

    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.right_margin=Cm(2.0); sec.left_margin=Cm(2.0)
    doc.styles['Normal'].font.name=FONT
    doc.styles['Normal'].font.size=Pt(13)

    # الغلاف
    para(doc,sb=20,sa=0)
    t=doc.add_table(1,1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,'1F497D')
    p=c.paragraphs[0]; set_rtl(p); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
    run(p,"وثيقة مفهوم المشروع",18,True,C_WHITE)
    p2=c.add_paragraph(); set_rtl(p2); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after=Pt(14)
    run(p2,PROJECT["العنوان"],20,True,C_WHITE)

    t2=doc.add_table(1,1); t2.alignment=WD_TABLE_ALIGNMENT.CENTER
    c2=t2.cell(0,0); shade(c2,'DEF0FA')
    p3=c2.paragraphs[0]; set_rtl(p3); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before=Pt(10); p3.paragraph_format.space_after=Pt(10)
    run(p3,f"صندوق المناخ الأخضر GCF  |  {PROJECT['الدولة']}  |  {today}",13,False,C_DARK)
    para(doc,sb=8,sa=4)

    # بطاقات ملخص
    kpis=[
        ("تمويل GCF",       f"${PROJECT['تمويل_GCF_مليون']}M", "1F497D"),
        ("تمويل مشترك",     f"${PROJECT['تمويل_مشترك_مليون']}M","37864B"),
        ("إجمالي المشروع",  f"${total}M",                      "2E75B6"),
    ]
    tbl_k=doc.add_table(2,3); tbl_k.style="Table Grid"; tbl_k.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(l,v,bg) in enumerate(kpis):
        ctxt(tbl_k.cell(0,i),l,12,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl_k.cell(1,i),v,18,True,RGBColor(*bytes.fromhex(bg)),WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # القسم أ: معلومات المشروع
    heading(doc,"القسم (أ): معلومات المشروع الأساسية",1)
    info_table(doc,[
        ("عنوان المشروع",     PROJECT["العنوان"]),
        ("الدولة",            PROJECT["الدولة"]),
        ("الجهة الوطنية NDA", PROJECT["NDA"]),
        ("الكيان المعتمد",    PROJECT["الكيان_المعتمد"]),
        ("حجم المشروع",       PROJECT["حجم_المشروع"]),
        ("مجال النتائج",      PROJECT["مجال_النتائج"]),
        ("مدة التنفيذ",       f"{PROJECT['المدة_سنوات']} سنوات"),
        ("المستفيدون",        PROJECT["المستفيدون"]),
    ])

    # القسم ب: السياق والمشكلة
    heading(doc,"القسم (ب): السياق والمشكلة",1)
    heading(doc,"ب.1  سياق الهشاشة المناخية",2)
    para(doc,PROJECT["سياق_الهشاشة"],13,sa=8)

    heading(doc,"ب.2  العوائق الرئيسية",2)
    for barrier in PROJECT["العوائق"]:
        p=doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after=Pt(4)
        run(p,"✗  "+barrier,12,False,C_TEXT)

    heading(doc,"ب.3  نظرية التغيير",2)
    para(doc,PROJECT["نظرية_التغيير"],13,sa=8)

    # القسم ج: المكوّنات والمخرجات
    heading(doc,"القسم (ج): المكوّنات والمخرجات",1)
    heading(doc,"ج.1  مكوّنات المشروع والميزانية",2)
    hdrs=["المكوّن","التكلفة (مليون $)","المحور"]
    wids=[8.5,3.5,3.0]
    tbl_c=doc.add_table(1+len(PROJECT["المكوّنات"]),3)
    tbl_c.style="Table Grid"; tbl_c.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(h,w) in enumerate(zip(hdrs,wids)):
        cc=tbl_c.cell(0,i); cc.width=Cm(w)
        ctxt(cc,h,12,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"2E75B6")
    total_comp=0
    for ri,comp in enumerate(PROJECT["المكوّنات"],1):
        bg="F2F7FD" if ri%2==0 else "FFFFFF"
        ctxt(tbl_c.cell(ri,0),comp["المكوّن"],12,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl_c.cell(ri,1),f"{comp['التكلفة_مليون']}",13,True,C_DARK,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl_c.cell(ri,2),comp["المحور"],12,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
        total_comp+=comp["التكلفة_مليون"]
    para(doc,sb=4,sa=4)

    heading(doc,"ج.2  المخرجات المتوقعة",2)
    for out in PROJECT["المخرجات"]:
        p=doc.add_paragraph(); set_rtl(p); p.paragraph_format.space_after=Pt(4)
        run(p,"✓  "+out,12,False,C_TEXT)

    # القسم د: مؤشرات الأداء
    heading(doc,"القسم (د): مؤشرات الأداء",1)
    hdrs2=["المؤشر","الخط الأساس","الهدف"]
    wids2=[8.0,3.5,4.0]
    tbl_i=doc.add_table(1+len(PROJECT["مؤشرات_الأداء"]),3)
    tbl_i.style="Table Grid"; tbl_i.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,(h,w) in enumerate(zip(hdrs2,wids2)):
        cc=tbl_i.cell(0,i); cc.width=Cm(w)
        ctxt(cc,h,12,True,C_WHITE,WD_ALIGN_PARAGRAPH.CENTER,"1F497D")
    for ri,ind in enumerate(PROJECT["مؤشرات_الأداء"],1):
        bg="F2F7FD" if ri%2==0 else "FFFFFF"
        ctxt(tbl_i.cell(ri,0),ind["المؤشر"],12,False,C_TEXT,WD_ALIGN_PARAGRAPH.RIGHT,bg)
        ctxt(tbl_i.cell(ri,1),ind["الخط_الأساس"],12,False,C_TEXT,WD_ALIGN_PARAGRAPH.CENTER,bg)
        ctxt(tbl_i.cell(ri,2),ind["الهدف"],12,True,C_GREEN,WD_ALIGN_PARAGRAPH.CENTER,bg)

    # القسم هـ: متطلبات GCF
    heading(doc,"القسم (هـ): متطلبات صندوق المناخ الأخضر",1)
    info_table(doc,[
        ("التحوّل النموذجي",   PROJECT["التحوّل_النموذجي"]),
        ("ملكية الدولة",       "يستند إلى NDC العراق وخطة NAP الوطنية"),
        ("اعتبارات النوع",     PROJECT["اعتبارات_النوع"]),
        ("الضمانات البيئية",   PROJECT["الضمانات_البيئية"]),
        ("إطار المتابعة",      PROJECT["إطار_المتابعة"]),
    ])

    footer=doc.sections[-1].footer.paragraphs[0]
    set_rtl(footer); footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run(footer,f"وثيقة مفهوم GCF - {PROJECT['الدولة']}  |  {today}",10,False,C_GRAY)

    return doc

if __name__=="__main__":
    print("جارٍ إنشاء وثيقة مفهوم المشروع GCF...")
    doc=build_concept_note()
    out="tools/gcf_prep/output_GCF_وثيقة_المفهوم.docx"
    doc.save(out)
    total=PROJECT["تمويل_GCF_مليون"]+PROJECT["تمويل_مشترك_مليون"]
    print(f"✓ تم إنشاء الوثيقة: {out}")
    print(f"  تمويل GCF: ${PROJECT['تمويل_GCF_مليون']}M  |  إجمالي: ${total}M")
