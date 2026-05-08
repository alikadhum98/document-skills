"""
أداة تتبع التمويل المناخي وتحليل فجوات GCF - العراق
التشغيل: python tools/finance_tracker/finance_tracker.py
المخرجات: تقرير Word + رسوم بيانية + JSON
"""

import sys, os, io, json, datetime
sys.path.insert(0, os.path.dirname(__file__))
from finance_data import (NDC_NEEDS, GEF_PROJECTS, GCF_PROJECTS,
                           ADAPTATION_FUND, BILATERAL_FINANCE, OPPORTUNITIES)

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import arabic_reshaper
from bidi.algorithm import get_display
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ========== إعداد ==========
plt.rcParams['font.family']        = 'DejaVu Sans'
plt.rcParams['axes.unicode_minus'] = False

def ar(t):
    return get_display(arabic_reshaper.reshape(str(t)))

C_DARK  = RGBColor(0x1F, 0x49, 0x7D)
C_MED   = RGBColor(0x2E, 0x75, 0xB6)
C_GREEN = RGBColor(0x37, 0x86, 0x44)
C_RED   = RGBColor(0xC0, 0x00, 0x00)
C_AMBER = RGBColor(0xBF, 0x8F, 0x00)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY  = RGBColor(0x70, 0x70, 0x70)
C_TEXT  = RGBColor(0x26, 0x26, 0x26)
FONT    = "Simplified Arabic"

STATUS_BG = {
    "نشط":           "37864B",
    "مكتمل":         "1F497D",
    "قيد المراجعة":  "2E75B6",
    "تحضيري":        "BF8F00",
    "ملغى":          "C00000",
}

PRIORITY_BG = {
    "عالية جداً": "C00000",
    "عالية":      "BF8F00",
    "متوسطة":     "2E75B6",
}

# ========== حسابات ==========
def compute_summary():
    gef_total = sum(p["تمويل_GEF_مليون"] for p in GEF_PROJECTS)
    gef_co    = sum(p["تمويل_مشترك_مليون"] for p in GEF_PROJECTS)
    gcf_total = sum(p["تمويل_GCF_مليون"] for p in GCF_PROJECTS)
    gcf_co    = sum(p["تمويل_مشترك_مليون"] for p in GCF_PROJECTS)
    bil_total = sum(p["المبلغ_مليون"] for p in BILATERAL_FINANCE)
    total_mob = (gef_total + gef_co + gcf_total + gcf_co + bil_total) / 1000
    need      = NDC_NEEDS["الإجمالي_مليار_دولار"]
    gap       = need - total_mob
    cov       = total_mob / need * 100

    return {
        "GEF_مباشر_مليون":     round(gef_total, 2),
        "GEF_مشترك_مليون":     round(gef_co, 2),
        "GCF_مباشر_مليون":     round(gcf_total, 2),
        "GCF_مشترك_مليون":     round(gcf_co, 2),
        "ثنائي_مليون":          round(bil_total, 2),
        "إجمالي_مُعبَّأ_مليار": round(total_mob, 3),
        "الاحتياج_مليار":       need,
        "الفجوة_مليار":         round(gap, 3),
        "نسبة_التغطية_pct":     round(cov, 3),
        "مشاريع_نشطة":
            sum(1 for p in GEF_PROJECTS + GCF_PROJECTS
                if p.get("الحالة") in ["نشط", "قيد المراجعة"]),
    }

# ========== الرسوم البيانية ==========
def fig_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150,
                bbox_inches='tight', facecolor='white')
    buf.seek(0)
    plt.close(fig)
    return buf

def chart_gap():
    """شريط أفقي يُظهر الفجوة التمويلية"""
    s    = compute_summary()
    need = s["الاحتياج_مليار"]
    mob  = s["إجمالي_مُعبَّأ_مليار"]
    gap  = s["الفجوة_مليار"]

    fig, ax = plt.subplots(figsize=(10, 3))
    ax.barh([ar("الاحتياج الكلي")], [need], color='#E8F0F7', height=0.5,
            edgecolor='#1F497D', linewidth=1.5)
    ax.barh([ar("الاحتياج الكلي")], [mob],  color='#2E75B6', height=0.5, alpha=0.9)
    ax.text(mob + 0.5, 0, ar(f"مُعبَّأ: {mob:.2f}B$"), va='center',
            fontsize=11, color='#2E75B6', fontweight='bold')
    ax.text(need - 1, 0, ar(f"فجوة: {gap:.1f}B$"), va='center',
            ha='right', fontsize=11, color='#C00000', fontweight='bold')
    ax.set_xlim(0, need * 1.08)
    ax.set_xlabel(ar("مليار دولار"), fontsize=11)
    ax.set_title(ar("الفجوة بين الاحتياجات والتمويل المُعبَّأ"), fontsize=13, fontweight='bold')
    ax.spines[['top', 'right', 'left']].set_visible(False)
    ax.yaxis.set_visible(False)
    ax.xaxis.set_major_formatter(
        matplotlib.ticker.FuncFormatter(lambda x, _: f'{x:.0f}B'))
    fig.tight_layout()
    return fig_bytes(fig)

def chart_sources():
    """دائري لمصادر التمويل"""
    s      = compute_summary()
    labels = [ar("GEF مباشر"), ar("GEF مشترك"),
              ar("GCF مباشر"), ar("GCF مشترك"), ar("ثنائي")]
    values = [s["GEF_مباشر_مليون"], s["GEF_مشترك_مليون"],
              s["GCF_مباشر_مليون"], s["GCF_مشترك_مليون"], s["ثنائي_مليون"]]
    colors = ['#1F497D','#4472C4','#2E75B6','#9DC3E6','#70AD47']
    fig, ax = plt.subplots(figsize=(7, 6))
    wedges, texts, autotexts = ax.pie(
        values, labels=labels, autopct='%1.1f%%',
        colors=colors, startangle=140, pctdistance=0.82,
        explode=[0.04]*5)
    for at in autotexts:
        at.set_fontsize(10)
        at.set_fontweight('bold')
    ax.set_title(ar("توزيع مصادر التمويل المُعبَّأ (مليون $)"),
                 fontsize=13, fontweight='bold')
    fig.tight_layout()
    return fig_bytes(fig)

def chart_sectors():
    """شريط أفقي لتوزيع التمويل حسب القطاع"""
    sector_map = {}
    for p in GEF_PROJECTS + GCF_PROJECTS:
        sec = p["القطاع"]
        amt = p.get("تمويل_GEF_مليون", 0) + p.get("تمويل_GCF_مليون", 0)
        sector_map[sec] = sector_map.get(sec, 0) + amt
    for p in BILATERAL_FINANCE:
        sec = p["القطاع"]
        sector_map[sec] = sector_map.get(sec, 0) + p["المبلغ_مليون"] * 0.1

    sectors = sorted(sector_map, key=sector_map.get)
    values  = [sector_map[s] for s in sectors]
    labels  = [ar(s) for s in sectors]
    colors  = ['#1F497D','#2E75B6','#4472C4','#9DC3E6',
               '#70AD47','#A9D18E','#BF8F00','#FFD966'][:len(sectors)]

    fig, ax = plt.subplots(figsize=(9, 5))
    bars = ax.barh(labels, values, color=colors, alpha=0.85, height=0.6)
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2,
                f'${val:.1f}M', va='center', fontsize=10)
    ax.set_xlabel(ar("مليون دولار"), fontsize=11)
    ax.set_title(ar("توزيع التمويل المناخي المباشر حسب القطاع"),
                 fontsize=13, fontweight='bold')
    ax.grid(axis='x', alpha=0.3)
    ax.spines[['top', 'right']].set_visible(False)
    fig.tight_layout()
    return fig_bytes(fig)

def chart_timeline():
    """خط زمني للمشاريع"""
    all_proj = GEF_PROJECTS + GCF_PROJECTS
    names    = [ar(p["العنوان"][:22]+"...") if len(p["العنوان"]) > 22
                else ar(p["العنوان"]) for p in all_proj]
    starts   = [p.get("سنة_الاعتماد", p.get("سنة_التقديم", 2020))
                for p in all_proj]
    ends     = [p.get("سنة_الانتهاء", starts[i]+3)
                for i, p in enumerate(all_proj)]
    colors   = ['#37864B' if p["الحالة"] == "نشط"
                else '#2E75B6' if p["الحالة"] in ["قيد المراجعة","تحضيري"]
                else '#9E9E9E' for p in all_proj]

    fig, ax = plt.subplots(figsize=(11, 5))
    for i, (name, s, e, c) in enumerate(zip(names, starts, ends, colors)):
        ax.barh(i, e-s, left=s, color=c, alpha=0.8, height=0.5)
        ax.text(s - 0.1, i, name, ha='right', va='center', fontsize=9)

    ax.set_xlim(2017, 2031)
    ax.set_xlabel(ar("السنة"), fontsize=11)
    ax.set_title(ar("الجدول الزمني لمشاريع التمويل المناخي"), fontsize=13, fontweight='bold')
    legend = [
        mpatches.Patch(color='#37864B', label=ar("نشط")),
        mpatches.Patch(color='#2E75B6', label=ar("مراجعة/تحضير")),
        mpatches.Patch(color='#9E9E9E', label=ar("مكتمل")),
    ]
    ax.legend(handles=legend, fontsize=10, loc='lower right')
    ax.grid(axis='x', alpha=0.3)
    ax.yaxis.set_visible(False)
    ax.spines[['top','right','left']].set_visible(False)
    fig.tight_layout()
    return fig_bytes(fig)

# ========== دوال Word مساعدة ==========
def set_rtl(p):
    p._p.get_or_add_pPr().insert(0, OxmlElement('w:bidi'))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def run(p, txt, size=13, bold=False, color=None):
    r = p.add_run(txt)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color or C_TEXT
    return r

def para(doc, txt="", size=13, bold=False, color=None, sb=0, sa=5):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if txt:
        run(p, txt, size, bold, color)
    return p

def heading(doc, txt, level=1):
    c = {1: C_DARK, 2: C_MED}
    s = {1: 18, 2: 15}
    p = para(doc, txt, s[level], True, c[level], sb=10, sa=4)
    if level == 1:
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:color'), '1F497D')
        pBdr.append(bot)
        p._p.get_or_add_pPr().append(pBdr)
    return p

def shade(cell, hex_c):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_c)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)

def ctxt(cell, txt, size=12, bold=False, color=None,
         align=WD_ALIGN_PARAGRAPH.RIGHT, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_rtl(p)
    p.alignment = align
    run(p, txt, size, bold, color or C_TEXT)
    if bg:
        shade(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def insert_chart(doc, buf, caption, width=15):
    doc.add_picture(buf, width=Cm(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = para(doc, caption, 11, False, C_GRAY, sb=2, sa=8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ========== بناء التقرير ==========
def build_report():
    doc   = Document()
    today = datetime.date.today().strftime("%d/%m/%Y")
    s_    = compute_summary()

    for sec in doc.sections:
        sec.page_width    = Cm(21)
        sec.page_height   = Cm(29.7)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.right_margin  = Cm(2.0)
        sec.left_margin   = Cm(2.0)
    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(13)

    # === الغلاف ===
    para(doc, sb=20, sa=0)
    t = doc.add_table(1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, '1F497D')
    p = c.paragraphs[0]
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    run(p, "تقرير تتبع التمويل المناخي", 22, True, C_WHITE)
    p2 = c.add_paragraph()
    set_rtl(p2)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(16)
    run(p2, "تحليل الفجوات وفرص GCF - العراق", 15, False, C_WHITE)
    para(doc, f"تاريخ التقرير: {today}  |  وزارة البيئة - قسم التمويل المناخي",
         12, False, C_GRAY, sb=8, sa=4)
    doc.add_page_break()

    # === القسم الأول: الملخص التنفيذي ===
    heading(doc, "أولاً: الملخص التنفيذي", 1)

    kpis = [
        ("الاحتياج الكلي 2021-2030", f"{s_['الاحتياج_مليار']} مليار $",      "C00000"),
        ("التمويل المُعبَّأ",         f"{s_['إجمالي_مُعبَّأ_مليار']:.2f} مليار $", "2E75B6"),
        ("الفجوة التمويلية",          f"{s_['الفجوة_مليار']:.1f} مليار $",    "C00000"),
        ("نسبة التغطية",              f"{s_['نسبة_التغطية_pct']:.2f}%",       "BF8F00"),
        ("مشاريع نشطة",              str(s_["مشاريع_نشطة"]),                  "37864B"),
        ("فرص غير مستغلة",           str(len(OPPORTUNITIES)),                  "7030A0"),
    ]
    tbl_k = doc.add_table(2, 3)
    tbl_k.style = "Table Grid"
    tbl_k.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(3):
        ctxt(tbl_k.cell(0, i), kpis[i][0], 11, True, C_WHITE,
             WD_ALIGN_PARAGRAPH.CENTER, kpis[i][2])
        ctxt(tbl_k.cell(1, i), kpis[i][1], 15, True,
             RGBColor(*bytes.fromhex(kpis[i][2])),
             WD_ALIGN_PARAGRAPH.CENTER)

    para(doc, sb=5, sa=0)
    tbl_k2 = doc.add_table(2, 3)
    tbl_k2.style = "Table Grid"
    tbl_k2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i in range(3):
        ctxt(tbl_k2.cell(0, i), kpis[i+3][0], 11, True, C_WHITE,
             WD_ALIGN_PARAGRAPH.CENTER, kpis[i+3][2])
        ctxt(tbl_k2.cell(1, i), kpis[i+3][1], 15, True,
             RGBColor(*bytes.fromhex(kpis[i+3][2])),
             WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # === القسم الثاني: الفجوة التمويلية ===
    heading(doc, "ثانياً: تحليل الفجوة التمويلية", 1)
    para(doc,
         f"تحتاج المساهمة الوطنية المحددة للعراق إلى {s_['الاحتياج_مليار']} مليار دولار "
         f"خلال الفترة {NDC_NEEDS['الفترة']}. تُغطّي الموارد المُعبَّأ حتى الآن "
         f"{s_['إجمالي_مُعبَّأ_مليار']:.2f} مليار دولار فقط، "
         f"أي ما يعادل {s_['نسبة_التغطية_pct']:.2f}% من الاحتياج الكلي.",
         13, sa=8)
    insert_chart(doc, chart_gap(), "شكل (1): الفجوة بين الاحتياجات والتمويل المُعبَّأ")
    insert_chart(doc, chart_sources(), "شكل (2): توزيع مصادر التمويل المُعبَّأ", width=12)

    # === القسم الثالث: مشاريع GEF ===
    heading(doc, "ثالثاً: مشاريع صندوق البيئة العالمي GEF", 1)
    hdrs = ["المعرّف", "العنوان", "الحالة", "تمويل GEF (م$)", "تمويل مشترك (م$)", "الوكالة"]
    wids = [2.5, 5.5, 2.5, 2.8, 3.0, 2.0]
    tbl_g = doc.add_table(1 + len(GEF_PROJECTS), 6)
    tbl_g.style = "Table Grid"
    tbl_g.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (h, w) in enumerate(zip(hdrs, wids)):
        c = tbl_g.cell(0, i)
        c.width = Cm(w)
        ctxt(c, h, 11, True, C_WHITE, WD_ALIGN_PARAGRAPH.CENTER, '2E75B6')
    for ri, p_ in enumerate(GEF_PROJECTS, 1):
        bg  = 'F2F7FD' if ri % 2 == 0 else 'FFFFFF'
        sbg = STATUS_BG.get(p_["الحالة"], "888888")
        ctxt(tbl_g.cell(ri, 0), p_["المعرّف"],  11, False, C_TEXT, WD_ALIGN_PARAGRAPH.CENTER, bg)
        ctxt(tbl_g.cell(ri, 1), p_["العنوان"],  11, False, C_TEXT, WD_ALIGN_PARAGRAPH.RIGHT,  bg)
        ctxt(tbl_g.cell(ri, 2), p_["الحالة"],   11, True,  C_WHITE,WD_ALIGN_PARAGRAPH.CENTER, sbg)
        ctxt(tbl_g.cell(ri, 3), f"{p_['تمويل_GEF_مليون']:.2f}", 12, True,  C_DARK, WD_ALIGN_PARAGRAPH.CENTER, bg)
        ctxt(tbl_g.cell(ri, 4), f"{p_['تمويل_مشترك_مليون']:.2f}", 11, False, C_TEXT, WD_ALIGN_PARAGRAPH.CENTER, bg)
        ctxt(tbl_g.cell(ri, 5), p_["الوكالة"],  11, False, C_TEXT, WD_ALIGN_PARAGRAPH.CENTER, bg)

    para(doc, sb=8, sa=4)

    # === القسم الرابع: مشاريع GCF ===
    heading(doc, "رابعاً: مشاريع صندوق المناخ الأخضر GCF", 1)
    for p_ in GCF_PROJECTS:
        heading(doc, f"{p_['المعرّف']} - {p_['العنوان'][:40]}", 2)
        sbg = STATUS_BG.get(p_["الحالة"], "888888")
        rows = [
            ("الحالة",              p_["الحالة"],                        sbg),
            ("القطاع / المحور",     f"{p_['القطاع']} | {p_['المحور']}", None),
            ("تمويل GCF",           f"{p_['تمويل_GCF_مليون']} مليون $", None),
            ("تمويل مشترك",         f"{p_['تمويل_مشترك_مليون']} مليون $", None),
            ("إجمالي المشروع",      f"{p_['تمويل_GCF_مليون']+p_['تمويل_مشترك_مليون']:.1f} مليون $", None),
            ("الكيان المعتمد",      p_["الكيان_المعتمد"],               None),
            ("المستفيدون",          p_["المستفيدون"],                    None),
            ("مرحلة الدورة",        p_["مرحلة_الدورة"],                 None),
        ]
        tbl_p = doc.add_table(len(rows), 2)
        tbl_p.style = "Table Grid"
        for ri, (label, val, bg_) in enumerate(rows):
            bg = bg_ if bg_ else ("EEF4FB" if ri % 2 == 0 else "FFFFFF")
            vc = C_WHITE if bg_ else C_TEXT
            vb = bool(bg_)
            ctxt(tbl_p.cell(ri, 0), label, 12, True,  C_DARK, WD_ALIGN_PARAGRAPH.RIGHT, "EEF4FB" if not bg_ else bg)
            ctxt(tbl_p.cell(ri, 1), val,   12, vb,    vc,     WD_ALIGN_PARAGRAPH.RIGHT, bg)
        para(doc, sb=6, sa=4)

    doc.add_page_break()

    # === القسم الخامس: الرسوم =====
    heading(doc, "خامساً: التحليل البياني", 1)
    insert_chart(doc, chart_sectors(), "شكل (3): توزيع التمويل المناخي حسب القطاع")
    insert_chart(doc, chart_timeline(), "شكل (4): الجدول الزمني لمشاريع التمويل المناخي")

    doc.add_page_break()

    # === القسم السادس: الفرص غير المستغلة ===
    heading(doc, "سادساً: فرص التمويل غير المستغلة", 1)
    para(doc, "تُحدِّد الجداول التالية أبرز مصادر التمويل المتاحة التي لم يُقدَّم إليها "
         "بعد، مرتّبةً حسب الأولوية وسهولة الوصول.", 13, sa=8)

    hdrs2 = ["المصدر", "المبلغ المتاح", "النوع", "الأولوية", "ملاحظة"]
    wids2 = [3.5, 3.5, 3.0, 2.5, 5.0]
    tbl_o = doc.add_table(1 + len(OPPORTUNITIES), 5)
    tbl_o.style = "Table Grid"
    tbl_o.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (h, w) in enumerate(zip(hdrs2, wids2)):
        c = tbl_o.cell(0, i)
        c.width = Cm(w)
        ctxt(c, h, 11, True, C_WHITE, WD_ALIGN_PARAGRAPH.CENTER, '1F497D')
    for ri, op in enumerate(OPPORTUNITIES, 1):
        bg   = 'F2F7FD' if ri % 2 == 0 else 'FFFFFF'
        pbg  = PRIORITY_BG.get(op["الأولوية"], "888888")
        ctxt(tbl_o.cell(ri, 0), op["المصدر"],       12, True,  C_DARK,  WD_ALIGN_PARAGRAPH.RIGHT,  bg)
        ctxt(tbl_o.cell(ri, 1), op["المبلغ_المتاح"],11, False, C_TEXT,  WD_ALIGN_PARAGRAPH.CENTER, bg)
        ctxt(tbl_o.cell(ri, 2), op["النوع"],         11, False, C_TEXT,  WD_ALIGN_PARAGRAPH.RIGHT,  bg)
        ctxt(tbl_o.cell(ri, 3), op["الأولوية"],      11, True,  C_WHITE, WD_ALIGN_PARAGRAPH.CENTER, pbg)
        ctxt(tbl_o.cell(ri, 4), op["ملاحظة"],        11, False, C_GRAY,  WD_ALIGN_PARAGRAPH.RIGHT,  bg)

    # === التوصيات ===
    heading(doc, "سابعاً: التوصيات الأولوية", 1)
    recs = [
        "تقديم طلب فوري لصندوق التكيف - لم يُقدَّم أي طلب حتى الآن وهو أسرع المسارات",
        "تفعيل دعم الاستعداد GCF Readiness لتحضير مشاريع جديدة بتمويل حتى 3 مليون $ / سنة",
        "متابعة مشروع المياه في GCF والتأكد من اجتياز مرحلة المراجعة",
        "تقديم مشروع طاقة شمسية ريفية إلى GCF بعد اكتمال التحضير",
        "التقدّم لدورة LDCF/SCCF القادمة بمشروع تكيف زراعي",
        "الاستعانة بتمويل GEF-8 لمشروع إدارة المياه الجوفية",
    ]
    for i, rec in enumerate(recs, 1):
        p = doc.add_paragraph()
        set_rtl(p)
        p.paragraph_format.space_after = Pt(4)
        run(p, f"{i}.  ", 13, True, C_MED)
        run(p, rec, 13, False, C_TEXT)

    # === التذييل ===
    footer = doc.sections[-1].footer.paragraphs[0]
    set_rtl(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(footer,
        f"تقرير التمويل المناخي - العراق  |  {today}  |  وزارة البيئة",
        10, False, C_GRAY)

    return doc, s_

# ========== التشغيل ==========
if __name__ == "__main__":
    print("جارٍ إنشاء تقرير التمويل المناخي...")
    doc, summary = build_report()

    out_docx = "tools/finance_tracker/output_Finance_تقرير_التمويل.docx"
    doc.save(out_docx)
    print(f"✓ تم إنشاء التقرير: {out_docx}")

    out_json = "tools/finance_tracker/output_Finance_ملخص.json"
    with open(out_json, 'w', encoding='utf-8') as f:
        json.dump({
            "تاريخ_التقرير": str(datetime.date.today()),
            "ملخص": summary,
            "GEF": [{"العنوان": p["العنوان"], "الحالة": p["الحالة"],
                     "تمويل_مليون": p["تمويل_GEF_مليون"]} for p in GEF_PROJECTS],
            "GCF": [{"العنوان": p["العنوان"], "الحالة": p["الحالة"],
                     "تمويل_مليون": p["تمويل_GCF_مليون"]} for p in GCF_PROJECTS],
        }, f, ensure_ascii=False, indent=2)
    print(f"✓ تم حفظ الملخص: {out_json}")

    print("\n" + "="*55)
    print("الملخص التنفيذي")
    print("="*55)
    print(f"  الاحتياج الكلي:      {summary['الاحتياج_مليار']} مليار $")
    print(f"  التمويل المُعبَّأ:    {summary['إجمالي_مُعبَّأ_مليار']:.3f} مليار $")
    print(f"  الفجوة التمويلية:    {summary['الفجوة_مليار']:.2f} مليار $")
    print(f"  نسبة التغطية:        {summary['نسبة_التغطية_pct']:.3f}%")
    print(f"  فرص غير مستغلة:     {len(OPPORTUNITIES)}")
